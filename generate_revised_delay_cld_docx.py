"""
Generate a revised causal loop diagram for the EV supply-chain SD layer,
explicitly representing system delays and grounding them in available model data.

Output:
  ABM_SD_Revised_Causal_Loop_with_Delays.docx
  results/sd_revised_causal_loop_with_delays.png
"""

from __future__ import annotations

from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from model.config import CELL_MAKERS, MINERALS, OEMS, TIER1
from model.sd_model import TARGET_WEEKS


ROOT = Path(__file__).resolve().parent
RESULTS_DIR = ROOT / "results"
OUTPUT_DOCX = ROOT / "ABM_SD_Revised_Causal_Loop_with_Delays.docx"
OUTPUT_CLD = RESULTS_DIR / "sd_revised_causal_loop_with_delays.png"


NODES = {
    "EV demand": (0.50, 0.94),
    "OEM backlog": (0.79, 0.91),
    "OEM production": (0.74, 0.72),
    "Component demand": (0.88, 0.50),
    "Tier-1 order pipeline": (0.76, 0.31),
    "Tier-1 production": (0.57, 0.20),
    "Component inventory": (0.35, 0.22),
    "Cell demand": (0.20, 0.36),
    "Cell production": (0.15, 0.60),
    "Mineral demand": (0.25, 0.80),
    "Mineral stocks": (0.47, 0.74),
    "Input availability": (0.49, 0.50),
    "Price signal": (0.28, 0.52),
    "Capacity investment": (0.55, 0.06),
    "Effective capacity": (0.73, 0.10),
}

# source, target, polarity, delay, interpretation
EDGES = [
    ("EV demand", "OEM production", "+", "short", "Higher demand raises desired vehicle output."),
    ("EV demand", "OEM backlog", "+", "short", "Demand not immediately fulfilled accumulates as backlog."),
    ("OEM production", "OEM backlog", "-", "short", "Production fulfils orders and reduces backlog."),
    ("OEM backlog", "OEM production", "+", "medium", "Backlog creates catch-up pressure, limited by component availability."),
    ("OEM production", "Component demand", "+", "short", "Each vehicle requires packs, inverters, motors, and harnesses."),
    ("Component demand", "Component inventory", "-", "short", "Component use draws down inventory."),
    ("Component inventory", "OEM production", "+", "short", "Higher component inventory enables more vehicle output."),
    ("Component inventory", "Tier-1 order pipeline", "-", "short", "Low inventory triggers replenishment orders and bullwhip amplification."),
    ("Tier-1 order pipeline", "Tier-1 production", "+", "delayed", "Orders translate into production and deliveries after lead times."),
    ("Tier-1 production", "Component inventory", "+", "delayed", "Subsystem output replenishes component inventory after production/logistics delay."),
    ("Component demand", "Cell demand", "+", "short", "Battery pack demand pulls cell demand."),
    ("Cell demand", "Cell production", "+", "short", "Cell makers respond to demand subject to minerals and capacity."),
    ("Cell production", "Component inventory", "+", "medium", "Cell output supports pack availability."),
    ("Cell production", "Mineral demand", "+", "short", "Cell production consumes lithium, cobalt, graphite, and related materials."),
    ("Mineral demand", "Mineral stocks", "-", "short", "Material consumption draws down mineral stocks."),
    ("Mineral stocks", "Input availability", "+", "short", "Higher stocks improve input availability fractions."),
    ("Input availability", "Cell production", "+", "short", "Low mineral availability constrains cell output."),
    ("Input availability", "Tier-1 production", "+", "short", "Low REE or SiC availability constrains motor and inverter output."),
    ("Mineral stocks", "Price signal", "-", "short", "Lower mineral stocks increase scarcity price pressure."),
    ("Price signal", "EV demand", "-", "medium", "Higher prices reduce demand through elasticity with market response delay."),
    ("OEM backlog", "Capacity investment", "+", "delayed", "Persistent backlog encourages capacity expansion."),
    ("Price signal", "Capacity investment", "-", "delayed", "High input prices discourage or defer investment."),
    ("Capacity investment", "Effective capacity", "+", "long", "Investment becomes usable capacity only after construction/ramp-up."),
    ("Effective capacity", "Cell production", "+", "long", "More effective capacity lifts output potential."),
    ("Effective capacity", "Tier-1 production", "+", "long", "More effective capacity lifts subsystem output potential."),
]

LOOPS = [
    [
        "B1 - Scarcity price balancing loop",
        "Mineral demand depletes mineral stocks; lower stocks raise the price signal; higher prices reduce EV demand; lower demand reduces mineral demand.",
        "Balancing",
        "Demand response is not instantaneous; price pass-through and purchasing decisions create a medium delay.",
    ],
    [
        "B2 - Inventory replenishment loop",
        "Component demand reduces component inventory; low inventory triggers orders; orders move through the Tier-1 pipeline; production replenishes inventory.",
        "Balancing with delay",
        "Lead-time delays can create overshoot and oscillation, especially with bullwhip ordering.",
    ],
    [
        "R1 - Backlog pressure loop",
        "EV demand increases backlog; backlog raises desired OEM production; production reduces backlog only if components are available.",
        "Reinforcing until constrained",
        "The loop is constrained by component inventory and input availability.",
    ],
    [
        "R2 - Capacity expansion loop",
        "Backlog and demand encourage investment; investment increases effective capacity after a long delay; capacity supports higher production.",
        "Reinforcing with long delay",
        "Long construction and qualification delays mean the loop cannot solve short-run shocks.",
    ],
    [
        "B3 - Financial/investment discipline loop",
        "Price pressure and weak margins reduce investment willingness; lower investment slows capacity growth and keeps bottlenecks binding.",
        "Balancing / constraint loop",
        "This loop links the new listed-company financial calibration to SD capacity dynamics.",
    ],
]


def delay_table_rows() -> list[list[object]]:
    rows = []
    for name, cfg in TIER1.items():
        rows.append(
            [
                f"Tier-1 {name}",
                cfg["lead_time_weeks"],
                cfg["safety_stock_weeks"],
                "Order pipeline and subsystem replenishment delay.",
            ]
        )
    for mineral, cfg in MINERALS.items():
        rows.append(
            [
                f"Mineral {mineral}",
                "shock recovery dependent",
                cfg["safety_stock_weeks"],
                "Inventory buffer delays downstream shortage propagation.",
            ]
        )
    rows.append(["Cell makers", "weekly capacity growth, shock recovery", "4-6 typical", "Capacity and inventory response varies by firm financial profile."])
    rows.append(["OEM groups", "weekly backlog catch-up", "3-6 typical", "Backlog creates pressure but production is component constrained."])
    return rows


def draw_diagram() -> None:
    RESULTS_DIR.mkdir(exist_ok=True)
    fig, ax = plt.subplots(figsize=(14, 9))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    delay_style = {
        "short": ("-", 0.82),
        "medium": ("--", 0.78),
        "delayed": ((0, (2, 2)), 0.76),
        "long": ((0, (5, 3)), 0.72),
    }

    for source, target, polarity, delay, label in EDGES:
        x1, y1 = NODES[source]
        x2, y2 = NODES[target]
        color = "#2563eb" if polarity == "+" else "#dc2626"
        linestyle, alpha = delay_style.get(delay, ("-", 0.8))
        arrow = FancyArrowPatch(
            (x1, y1),
            (x2, y2),
            arrowstyle="-|>",
            mutation_scale=14,
            linewidth=1.25,
            color=color,
            linestyle=linestyle,
            connectionstyle="arc3,rad=0.08",
            alpha=alpha,
        )
        ax.add_patch(arrow)
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        text = polarity if delay == "short" else f"{polarity} ||"
        ax.text(
            mx,
            my,
            text,
            fontsize=9,
            fontweight="bold",
            color=color,
            bbox=dict(boxstyle="round,pad=0.18", fc="white", ec=color, lw=0.7),
            ha="center",
            va="center",
        )

    for node, (x, y) in NODES.items():
        ax.text(
            x,
            y,
            node,
            ha="center",
            va="center",
            fontsize=9.5,
            fontweight="bold",
            bbox=dict(boxstyle="round,pad=0.42", fc="#f8fafc", ec="#334155", lw=1.1),
        )

    ax.text(0.02, 0.065, "+ same-direction effect", fontsize=9, color="#2563eb")
    ax.text(0.02, 0.038, "- opposite-direction effect", fontsize=9, color="#dc2626")
    ax.text(0.02, 0.011, "|| indicates material delay; dashed arrows indicate medium/long delays", fontsize=9, color="#334155")
    ax.set_title("Revised SD Causal Loop Diagram with Explicit Delays", fontsize=14, fontweight="bold")
    plt.tight_layout()
    plt.savefig(OUTPUT_CLD, dpi=180, bbox_inches="tight")
    plt.close(fig)


def add_table(document: Document, headers: list[str], rows: list[list[object]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        for paragraph in table.rows[0].cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = "" if value is None else str(value)
    document.add_paragraph()


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def build_document() -> Document:
    draw_diagram()
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Revised Causal Loop Diagram with Delays\nEV Supply Chain ABM + SD Model")
    run.bold = True
    run.font.size = Pt(18)

    document.add_heading("1. Analysis of the Previous Causal Loop Diagram", level=1)
    document.add_paragraph(
        "The earlier causal loop diagram correctly captured the main reinforcing and balancing feedbacks: demand pull, "
        "inventory replenishment, mineral scarcity pricing, and backlog catch-up. However, it was too instantaneous. "
        "The available model data show that delays are not secondary details; they are central mechanisms that explain "
        "overshoot, backlog persistence, and delayed recovery after shocks."
    )
    add_bullets(
        document,
        [
            "Tier-1 lead times differ substantially: battery packs use a 4-week lead time, harnesses 6 weeks, motors 12 weeks, and inverters 16 weeks.",
            "Safety stocks also differ: harness inventory is lean at 2 weeks, while SiC-related inverter buffers are much larger.",
            "Financial calibration changes recovery speed, growth, buffer capacity, and shock absorption by agent.",
            "Capacity expansion is structurally delayed because investment, construction, qualification, and ramp-up cannot occur within one weekly step.",
            "Price signals reduce demand with a lag because commodity prices must pass through component costs, vehicle prices, and consumer purchasing decisions.",
        ],
    )

    document.add_heading("2. Revised Delay-Aware Causal Loop Diagram", level=1)
    document.add_paragraph(
        "The revised diagram adds explicit pipeline, capacity, and market-response delays. The symbol || marks delayed causal effects."
    )
    document.add_picture(str(OUTPUT_CLD), width=Inches(7.4))

    document.add_heading("3. Revised Feedback Loop Interpretation", level=1)
    add_table(document, ["Loop", "Description", "Type", "Delay interpretation"], LOOPS)

    document.add_heading("4. Revised Causal Link Register", level=1)
    add_table(
        document,
        ["From", "To", "Polarity", "Delay", "Interpretation"],
        [[source, target, polarity, delay, label] for source, target, polarity, delay, label in EDGES],
    )

    document.add_heading("5. Delay Register Based on Available Model Data", level=1)
    add_table(
        document,
        ["Model element", "Delay parameter", "Buffer / safety stock", "How it affects system behaviour"],
        delay_table_rows(),
    )

    document.add_heading("6. Modelling Implications", level=1)
    add_bullets(
        document,
        [
            "The harness loop is fast to disrupt because safety stock is low, even though nominal lead time is shorter than inverters.",
            "The inverter loop is slower but more persistent because SiC supply and 16-week lead time create delayed replenishment.",
            "Backlog is not automatically stabilizing. It can reinforce demand pressure until component stocks or price signals constrain production.",
            "Financially stronger firms should recover faster and absorb shocks, but this does not remove physical lead-time delays.",
            "Capacity investment should be modelled as a delayed stock or pipeline rather than an immediate increase in weekly capacity.",
        ],
    )

    document.add_heading("7. Recommended Model Revisions", level=1)
    add_bullets(
        document,
        [
            "Add explicit capacity-investment pipeline stocks for cells, inverters, motors, and OEM assembly.",
            "Separate order placement from delivery for cell suppliers, not only Tier-1 suppliers.",
            "Report effective delay by subsystem in every shock scenario.",
            "Use financial profiles to affect recovery and investment rates, while keeping physical lead times as independent constraints.",
            "Track backlog age, not only backlog size, to distinguish temporary demand accumulation from persistent unmet demand.",
        ],
    )

    document.add_heading("8. Current Model Scale", level=1)
    add_table(
        document,
        ["Model object", "Count"],
        [
            ["Mineral groups", len(MINERALS)],
            ["Cell makers", len(CELL_MAKERS)],
            ["Tier-1 subsystem groups", len(TIER1)],
            ["OEM groups", len(OEMS)],
        ],
    )

    return document


def main() -> None:
    document = build_document()
    document.save(OUTPUT_DOCX)
    print(f"Wrote {OUTPUT_DOCX}")
    print(f"Wrote {OUTPUT_CLD}")


if __name__ == "__main__":
    main()
