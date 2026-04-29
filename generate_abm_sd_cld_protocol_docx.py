"""
Generate standardized reporting protocols for the ABM layer and a causal loop
diagram for the SD layer of the EV supply-chain model.

Output:
  ABM_SD_Reporting_Protocols_and_Causal_Loop.docx
  results/sd_causal_loop_diagram.png
"""

from __future__ import annotations

from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from model.config import CELL_MAKERS, MARKETS, MINERALS, OEMS, TIER1
from model.shocks import SCENARIOS


ROOT = Path(__file__).resolve().parent
RESULTS_DIR = ROOT / "results"
OUTPUT_DOCX = ROOT / "ABM_SD_Reporting_Protocols_and_Causal_Loop.docx"
OUTPUT_CLD = RESULTS_DIR / "sd_causal_loop_diagram.png"


CLD_NODES = {
    "EV demand": (0.50, 0.92),
    "OEM production": (0.72, 0.74),
    "Component demand": (0.82, 0.50),
    "Tier-1 production": (0.68, 0.26),
    "Component inventory": (0.44, 0.18),
    "Cell demand": (0.22, 0.32),
    "Cell production": (0.17, 0.58),
    "Mineral demand": (0.25, 0.78),
    "Mineral stocks": (0.48, 0.70),
    "Input availability": (0.52, 0.48),
    "Price signal": (0.28, 0.48),
    "Backlog": (0.78, 0.90),
}

CLD_EDGES = [
    ("EV demand", "OEM production", "+", "demand pull"),
    ("OEM production", "Component demand", "+", "1:1 vehicle inputs"),
    ("Component demand", "Component inventory", "-", "drawdown"),
    ("Component inventory", "Tier-1 production", "+", "availability"),
    ("Tier-1 production", "Component inventory", "+", "replenishment"),
    ("OEM production", "Backlog", "-", "fulfilment"),
    ("EV demand", "Backlog", "+", "unmet demand"),
    ("Backlog", "OEM production", "+", "catch-up production"),
    ("Component demand", "Cell demand", "+", "pack demand"),
    ("Cell demand", "Cell production", "+", "demand pull"),
    ("Cell production", "Component inventory", "+", "pack supply"),
    ("Cell production", "Mineral demand", "+", "material consumption"),
    ("Mineral demand", "Mineral stocks", "-", "drawdown"),
    ("Mineral stocks", "Input availability", "+", "stock adequacy"),
    ("Input availability", "Cell production", "+", "Leontief constraint"),
    ("Input availability", "Tier-1 production", "+", "critical input constraint"),
    ("Mineral stocks", "Price signal", "-", "scarcity pricing"),
    ("Price signal", "EV demand", "-", "price elasticity"),
    ("Component inventory", "Tier-1 production", "-", "shortfall ordering/bullwhip"),
]

CLD_LOOPS = [
    [
        "B1 - Scarcity price balancing loop",
        "Mineral demand increases, mineral stocks fall, the price signal rises, EV demand weakens, and upstream material demand falls.",
        "Balancing",
    ],
    [
        "B2 - Inventory replenishment loop",
        "Component demand draws down inventory; lower inventory triggers higher Tier-1 production and ordering; replenishment rebuilds inventory.",
        "Balancing",
    ],
    [
        "R1 - Demand growth and production expansion loop",
        "EV demand raises OEM production, component demand, cell demand, and production activity, reinforcing capacity pressure in the supply chain.",
        "Reinforcing",
    ],
    [
        "R2 - Backlog catch-up loop",
        "Unmet EV demand increases backlog, which raises target OEM production in later periods until constrained by components.",
        "Reinforcing until constrained",
    ],
]


def add_title(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Standardized Reporting Protocols\nABM Model and SD Causal Loop Diagram")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("EV Supply Chain Hybrid Agent-Based Model and System Dynamics Framework")


def add_table(document: Document, headers: list[str], rows: list[list[object]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
        for paragraph in table.rows[0].cells[idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = "" if value is None else str(value)

    document.add_paragraph()


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Number")


def draw_causal_loop_diagram(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fig, ax = plt.subplots(figsize=(13, 8))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    for source, target, polarity, label in CLD_EDGES:
        x1, y1 = CLD_NODES[source]
        x2, y2 = CLD_NODES[target]
        color = "#1f77b4" if polarity == "+" else "#d62728"
        arrow = FancyArrowPatch(
            (x1, y1),
            (x2, y2),
            arrowstyle="-|>",
            mutation_scale=15,
            linewidth=1.35,
            color=color,
            connectionstyle="arc3,rad=0.08",
            alpha=0.82,
        )
        ax.add_patch(arrow)
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        ax.text(
            mx,
            my,
            polarity,
            fontsize=11,
            fontweight="bold",
            color=color,
            bbox=dict(boxstyle="circle,pad=0.18", fc="white", ec=color, lw=0.9),
            ha="center",
            va="center",
        )

    for node, (x, y) in CLD_NODES.items():
        ax.text(
            x,
            y,
            node,
            ha="center",
            va="center",
            fontsize=10,
            fontweight="bold",
            bbox=dict(boxstyle="round,pad=0.42", fc="#f8fafc", ec="#334155", lw=1.2),
        )

    ax.text(0.03, 0.04, "+ = same direction effect", fontsize=9, color="#1f77b4")
    ax.text(0.03, 0.01, "- = opposite direction effect", fontsize=9, color="#d62728")
    ax.set_title("System Dynamics Causal Loop Diagram - EV Supply Chain", fontsize=14, fontweight="bold")
    plt.tight_layout()
    plt.savefig(path, dpi=180, bbox_inches="tight")
    plt.close(fig)


def build_document() -> Document:
    draw_causal_loop_diagram(OUTPUT_CLD)

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)

    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    add_title(document)

    document.add_heading("1. Purpose and Reporting Boundary", level=1)
    document.add_paragraph(
        "This protocol standardizes how the EV supply-chain hybrid model should be reported. "
        "It covers the Agent-Based Model layer and the System Dynamics causal structure used "
        "to represent inventory pressure, production constraints, demand response, and shock propagation."
    )
    add_bullets(
        document,
        [
            "Model family: hybrid ABM + SD simulation.",
            "Time step: one week.",
            "Default horizon: 260 weeks.",
            "ABM entities: mineral sources, cell makers, Tier-1 subsystem suppliers, OEM groups, and market agents.",
            "SD entities: aggregate inventory stocks, availability fractions, price signal, demand response, and feedback loops.",
            "Reporting objective: reproducibility, scenario comparability, and transparent calibration.",
        ],
    )

    document.add_heading("2. Standardized Reporting Protocol for the ABM Model", level=1)
    document.add_heading("2.1 Minimum Metadata", level=2)
    add_table(
        document,
        ["Reporting item", "Required content"],
        [
            ["Model identifier", "EV Supply Chain ABM + SD Simulation."],
            ["Code version", "Git commit hash or archived code snapshot."],
            ["Run date", "Date and time the simulation was executed."],
            ["Random seed", "Default seed is 42 unless changed."],
            ["Scenario horizon", "Number of simulated weeks."],
            ["Scenario set", "All scenario IDs included in the run."],
            ["Financial calibration", "Whether listed-company financial profiles were enabled and which workbook was used."],
            ["Output files", "CSV, plot, and workbook paths produced by the run."],
        ],
    )

    document.add_heading("2.2 Agent Classes", level=2)
    add_table(
        document,
        ["Agent class", "Model role", "Key state variables", "Required reported outputs"],
        [
            [
                "MineralSupplierAgent",
                "Supplies lithium, cobalt, graphite, rare earths, and SiC source cohorts.",
                "mineral, country, global_share, shock_multiplier, recovery_rate_wk.",
                "weekly_supply_contribution, output_fraction, active shock status.",
            ],
            [
                "CellManufacturerAgent",
                "Converts mineral availability into battery cell production.",
                "weekly_capacity, market_share, LFP/NMC mix, inventory_gwh, backlog_gwh.",
                "output_gwh, inventory_gwh, backlog_gwh, utilisation.",
            ],
            [
                "Tier1SupplierAgent",
                "Produces packs, inverters, motors, and harnesses.",
                "weekly_capacity, inventory, pipeline, key_input, input_dependency, dual_source_active.",
                "output_k, inventory, shortage, dual-source activation.",
            ],
            [
                "OEMAgent",
                "Assembles vehicles using Leontief component inventories.",
                "weekly_target, component inventories, backlog_k, shock_multiplier.",
                "production_k, backlog_k, halt_weeks, average inventory.",
            ],
            [
                "MarketAgent",
                "Generates regional EV demand and price response.",
                "gwh_annual, yoy_growth, avg_kwh_per_vehicle, price_elasticity.",
                "weekly_demand_gwh, weekly_demand_k_vehicles.",
            ],
        ],
    )

    document.add_heading("2.3 Weekly ABM Execution Order", level=2)
    add_numbered(
        document,
        [
            "Apply scheduled shocks and shock resolutions.",
            "Compute SD input availability fractions from current stock levels.",
            "Step mineral supplier agents.",
            "Step cell manufacturer agents.",
            "Step Tier-1 subsystem supplier agents.",
            "Step OEM agents.",
            "Step market demand agents.",
            "Aggregate ABM outputs into SD inflow and outflow variables.",
            "Update SD stocks and record weekly metrics.",
        ],
    )

    document.add_heading("2.4 Agent Decision Rules by Tier", level=2)
    document.add_paragraph(
        "Agents are heterogeneous by tier and by subtype. Each agent makes weekly "
        "production, inventory, and capacity decisions from its local state plus "
        "shared SD signals for input availability, prices, and demand pressure."
    )
    add_table(
        document,
        ["Tier / agent type", "Production decision", "Inventory decision", "Capacity decision", "Main interaction rule"],
        [
            [
                "Tier 1 materials: lithium, cobalt, graphite, REE, SiC source agents",
                "Output equals source share multiplied by shock and recovery state. Source geography determines exposure, e.g. DRC cobalt or China graphite.",
                "No firm-level inventory ordering; material stocks are held in the SD layer as weeks of supply with transport delay.",
                "No endogenous firm capacity expansion; aggregate mineral supply scale grows in the SD layer by mineral-specific growth rates.",
                "Supplies flow into SD mineral stocks; downstream agents observe measured availability fractions.",
            ],
            [
                "Tier 2 cells: LFP-heavy makers such as BYD/CALB/CATL",
                "Produce up to capacity subject to lithium and graphite availability; cobalt constraint is weak because LFP share is high.",
                "Use an order-up-to rule for cell inventory; platform leaders replenish faster than hyper-scale challengers and NMC/NCA incumbents.",
                "Weekly capacity grows at the baseline cell growth rate, adjusted by financial growth multiplier.",
                "Serve market demand in proportion to cell-maker market share; output feeds pack production and mineral consumption.",
            ],
            [
                "Tier 2 cells: NMC/NCA-heavy makers such as LG ES, Panasonic, Samsung SDI, SK On, AESC",
                "Use the same cell production rule but with stronger cobalt exposure; high cobalt price reduces effective NMC dependence over time.",
                "Hold larger safety stocks for some makers; unmet demand becomes cell backlog.",
                "Capacity grows weekly and is financially modulated; SD cell capacity also has an investment pipeline.",
                "More sensitive to cobalt shocks than LFP-heavy agents, creating heterogeneous shock impacts.",
            ],
            [
                "Tier 3 battery-pack supplier",
                "Converts cell availability into pack output; production is tightly linked to the cells stock.",
                "Uses lead-time pipeline and order-up-to inventory policy with bullwhip amplification.",
                "Weekly capacity grows with demand trend and financial multiplier.",
                "Pack output is allocated to OEMs by OEM production share.",
            ],
            [
                "Tier 3 inverter supplier",
                "Produces inverters subject to partial SiC wafer dependency; non-SiC portion can continue under SiC shortage.",
                "Uses order-up-to inventory, long lead time, and dual sourcing when inventory falls below trigger.",
                "Capacity grows weekly but is constrained by specialised manufacturing and slower recovery.",
                "SiC scarcity lowers deliveries and propagates to OEM assembly.",
            ],
            [
                "Tier 3 motor supplier",
                "Produces motors subject to partial REE dependency; high REE prices trigger motor-design substitution that lowers effective REE dependence.",
                "Uses order-up-to inventory with pipeline deliveries and shortage tracking.",
                "Capacity grows weekly with financial adjustment.",
                "REE shortages constrain PMSM motor supply and can become an OEM bottleneck.",
            ],
            [
                "Tier 3 wiring-harness supplier",
                "Produces harnesses from capacity and shock state; currently no copper input dependency in ABM, while copper is tracked in SD for extension.",
                "Very low safety stock reflects JIT practice; shortages emerge quickly under disruption.",
                "Capacity grows weekly and recovery is relatively fast after shocks.",
                "Harness shocks propagate rapidly because every vehicle needs one harness set.",
            ],
            [
                "Tier 4 OEM groups",
                "Assemble vehicles using a Leontief rule over packs, inverters, motors, and harnesses; vertical integration cushions part of component shortages.",
                "Receive component deliveries, consume one of each component per vehicle, accumulate backlog from shortfalls, and clear backlog with surplus production.",
                "Weekly assembly target grows with the EV market and financial multiplier; shock recovery restores throughput gradually.",
                "OEM demand is allocated by annual production share; component shortages create production halts and backlog.",
            ],
            [
                "Market agents by region",
                "Do not produce physical goods; they set realised EV demand in GWh and vehicle units.",
                "No inventory stock; demand is reduced by high price signals and large industry backlog with region-specific sensitivity.",
                "No capacity decision; regional demand trend grows at configured YoY rate.",
                "Demand pulls cell production and OEM assembly targets; price and backlog feed back to demand.",
            ],
        ],
    )

    document.add_heading("2.5 Agent Calibration Reporting", level=2)
    add_table(
        document,
        ["Calibration field", "Required reporting rule"],
        [
            ["Capacity", "Report value, unit, source year, and transformation into weekly capacity."],
            ["Market share", "Report whether share is global, regional, or model-normalized."],
            ["Chemistry mix", "For cell makers, report LFP and NMC/NCA fractions."],
            ["Safety stock", "Report base safety stock weeks and any financial multiplier applied."],
            ["Recovery rate", "Report base recovery rate and financially adjusted recovery rate."],
            ["Shock absorption", "Report effective shock severity after financial shock absorption."],
            ["Supplier routing", "Report OEM-to-cell or OEM-to-component sourcing shares when used."],
        ],
    )

    document.add_heading("2.6 Scenario Reporting", level=2)
    add_table(
        document,
        ["Scenario field", "Definition"],
        [
            ["scenario_id", "Unique scenario key from model.shocks.SCENARIOS."],
            ["description", "Plain-language scenario description."],
            ["target_agent", "Agent receiving the disruption."],
            ["start_week / end_week", "Start and resolution week of the disruption."],
            ["nominal severity", "Fractional output loss before financial adjustment."],
            ["effective severity", "Fractional output loss after financial shock absorption."],
            ["peak loss", "Maximum production loss versus baseline."],
            ["recovery week", "Last week below threshold plus one."],
            ["cumulative loss", "Sum of weekly production shortfall."],
        ],
    )

    document.add_heading("3. System Dynamics Causal Loop Diagram", level=1)
    document.add_paragraph(
        "The SD layer represents feedback relationships among demand, production, inventories, input availability, "
        "and price response. The diagram below is a causal loop diagram, not a stock-flow diagram. "
        "It shows directional influence and polarity."
    )
    document.add_picture(str(OUTPUT_CLD), width=Inches(7.4))

    document.add_heading("3.1 Causal Link Register", level=2)
    add_table(
        document,
        ["From", "To", "Polarity", "Interpretation"],
        [[source, target, polarity, label] for source, target, polarity, label in CLD_EDGES],
    )

    document.add_heading("3.2 Feedback Loop Register", level=2)
    add_table(document, ["Loop", "Description", "Type"], CLD_LOOPS)

    document.add_heading("3.3 SD Variable Reporting Standard", level=2)
    add_table(
        document,
        ["Variable", "Reporting definition"],
        [
            ["EV demand", "Regional and aggregate battery demand from market agents."],
            ["OEM production", "Weekly vehicle output from OEM agents."],
            ["Component demand", "Vehicle-equivalent pack, inverter, motor, and harness requirements."],
            ["Tier-1 production", "Weekly subsystem output from Tier-1 agents."],
            ["Component inventory", "Vehicle-equivalent stocks of packs, inverters, motors, and harnesses."],
            ["Cell demand", "Cell GWh demanded by battery pack production."],
            ["Cell production", "Aggregate GWh output from cell maker agents."],
            ["Mineral demand", "Material consumption implied by cell, motor, and inverter production."],
            ["Mineral stocks", "Weeks of supply for lithium, cobalt, graphite, REE, and SiC wafer."],
            ["Input availability", "Stock divided by target stock, capped at 2.0."],
            ["Price signal", "Weighted scarcity index feeding demand elasticity."],
            ["Backlog", "Unmet vehicle demand carried forward by OEM agents."],
        ],
    )

    document.add_heading("4. Current Model Inventory for Reporting", level=1)
    add_table(
        document,
        ["Model object", "Count", "Code source"],
        [
            ["Mineral groups", len(MINERALS), "model.config.MINERALS"],
            ["Cell makers", len(CELL_MAKERS), "model.config.CELL_MAKERS"],
            ["Tier-1 subsystem groups", len(TIER1), "model.config.TIER1"],
            ["OEM groups", len(OEMS), "model.config.OEMS"],
            ["Markets", len(MARKETS), "model.config.MARKETS"],
            ["Scenario definitions", len(SCENARIOS), "model.shocks.SCENARIOS"],
        ],
    )

    document.add_heading("5. Reproducibility Checklist", level=1)
    add_numbered(
        document,
        [
            "Archive the exact model code and commit hash.",
            "Archive input Excel workbooks used for firm and financial calibration.",
            "Report package versions and Python version.",
            "Report seed, horizon, and scenario set.",
            "Export weekly time-series outputs for every scenario.",
            "Export scenario summary tables using a consistent baseline.",
            "Report causal loop assumptions and polarity definitions.",
            "Document all manual ticker, firm-name, or supplier-routing assumptions.",
        ],
    )

    return document


def main() -> None:
    document = build_document()
    document.save(OUTPUT_DOCX)
    print(f"Wrote {OUTPUT_DOCX}")
    print(f"Wrote {OUTPUT_CLD}")


if __name__ == "__main__":
    main()
