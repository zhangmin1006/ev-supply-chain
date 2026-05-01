"""Generate a methodology paper for the EV supply-chain ABM+SD model.

Run:
    python generate_methodology_paper_docx.py

Output:
    documents/EV_Supply_Chain_ABM_SD_Methodology_Paper.docx
"""

from __future__ import annotations

import os
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_PATH = os.path.join("documents", "EV_Supply_Chain_ABM_SD_Methodology_Paper.docx")

BLUE = RGBColor(31, 78, 121)
MID_BLUE = RGBColor(37, 99, 235)
DARK = RGBColor(15, 23, 42)
TEXT = RGBColor(51, 65, 85)
GREY = RGBColor(100, 116, 139)
LIGHT_BLUE = "DBEAFE"
LIGHT_GREY = "F8FAFC"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(section) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    add_page_number(section)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = TEXT

    for name, size, colour in [
        ("Title", 20, BLUE),
        ("Heading 1", 14, BLUE),
        ("Heading 2", 12, MID_BLUE),
        ("Heading 3", 11, DARK),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = colour


def heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True


def para(doc: Document, text: str = "", italic: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = TEXT
    r.italic = italic


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = TEXT


def eq(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.font.name = "Cambria Math"
    r.font.size = Pt(10.5)
    r.font.color.rgb = DARK


def caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(9)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GREY


def table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_repeat_table_header(t.rows[0])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        shade(c, "1F4E79")
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9)
    for row_idx, values in enumerate(rows):
        row = t.add_row()
        for i, value in enumerate(values):
            c = row.cells[i]
            if row_idx % 2 == 0:
                shade(c, LIGHT_GREY)
            c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = c.paragraphs[0]
            r = p.add_run(str(value))
            r.font.size = Pt(9)
            r.font.color.rgb = TEXT
    if widths:
        for i, width in enumerate(widths):
            for c in t.columns[i].cells:
                c.width = Cm(width)
    return t


def causal_loop_table(doc: Document) -> None:
    rows = [
        [
            "B1: Scarcity-price loop",
            "Mineral/component stock below target -> availability falls -> price index rises -> producers raise output or downstream demand weakens -> stock pressure eases.",
            "Balancing",
        ],
        [
            "B2: Chemistry substitution loop",
            "Cobalt price rises -> perceived cobalt price rises with delay -> LFP target share increases logistically -> cobalt demand per kWh falls -> cobalt stock recovers.",
            "Balancing",
        ],
        [
            "R1/B3: Capacity investment cycle",
            "High cell utilisation and favourable pack price margins trigger planning and construction of new capacity; new capacity later reduces utilisation pressure.",
            "Reinforcing then balancing",
        ],
        [
            "B4: Demand-adoption loop",
            "Vehicle price pressure and availability constraints reduce realised weekly EV demand relative to the exogenous adoption trend.",
            "Balancing",
        ],
        [
            "R2: Bullwhip amplification",
            "Shortfall at Tier-1 raises order-up-to quantities above demand; amplified orders lift upstream pressure and are smoothed into the bullwhip index.",
            "Reinforcing/tracking",
        ],
    ]
    table(doc, ["Loop", "Causal Logic", "Type"], rows, [3.4, 10.0, 2.8])
    caption(doc, "Table 4. Main causal-loop logic represented in the hybrid ABM-SD model.")


def build_document() -> None:
    os.makedirs("documents", exist_ok=True)
    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("A Hybrid Agent-Based and System Dynamics Model of EV Supply-Chain Resilience")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Methodology Paper")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = MID_BLUE

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(
        f"Queen's University Belfast | EV supply-chain simulation | Generated {date.today().isoformat()}"
    )
    r.font.size = Pt(10)
    r.font.color.rgb = GREY

    doc.add_paragraph()
    heading(doc, "Abstract", 1)
    para(
        doc,
        "This paper documents a hybrid Agent-Based Model (ABM) and System Dynamics (SD) model of electric-vehicle supply-chain resilience. "
        "The model represents critical mineral supply, battery-cell manufacturing, Tier-1 component supply, OEM assembly, and EV market demand over a weekly five-year horizon. "
        "ABM agents capture heterogeneous firm behaviour, including different sourcing structures, financial archetypes, inventory policies, chemistry mixes, and recovery rules. "
        "The SD layer tracks aggregate stocks, transport delays, measured inventory, price formation, chemistry substitution, capacity expansion, demand pressure, backlog, and bullwhip amplification. "
        "The paper describes the model purpose, scope, mathematics, logistics logic, ODD protocol, causal-loop structure, SD-ABM communication, assumptions, data sources, parameterisation, validation logic, and limitations."
    )

    heading(doc, "1. Purpose and Scope", 1)
    para(
        doc,
        "The model is designed to examine how disruption in upstream EV supply-chain nodes can propagate to downstream UK and global EV production. "
        "It is not a forecasting model in the narrow econometric sense. Instead, it is a structured simulation laboratory for comparing resilience mechanisms, shock pathways, policy packages, and operational assumptions."
    )
    bullet(doc, "Time horizon: 260 weekly steps, equivalent to five years.")
    bullet(doc, "Geographic focus: global supply-chain structure with explicit UK endpoint and separate Chinese, US, European, Korean, Japanese, and rest-of-world market/OEM groups.")
    bullet(doc, "Core outputs: weekly OEM production, UK OEM production, cell output, stock levels in weeks of supply, price indices, backlog, bullwhip index, and scenario loss metrics.")
    bullet(doc, "Core interventions: supply shocks, policy mitigation packages, and user-controlled parameter experiments in the Streamlit app.")

    heading(doc, "2. Model Architecture", 1)
    para(
        doc,
        "The model has a five-tier architecture. Material flows downstream from mineral supply to cells, components, vehicles, and markets. "
        "Information flows upstream through demand, backlog, price signals, availability fractions, and order-up-to behaviour."
    )
    table(
        doc,
        ["Tier", "Layer", "Agents or Stocks", "Main Output"],
        [
            ["0", "Critical minerals", "Lithium, cobalt, graphite, REE, SiC wafer, copper stocks; mineral supplier agents for all except copper", "Normalised weekly mineral supply and price indices"],
            ["1", "Battery cells", "CATL, BYD Cells, LG ES, Panasonic, Samsung SDI, SK On, CALB, AESC UK, others", "GWh of cell production and chemistry mix"],
            ["2", "Tier-1 components", "Battery pack, inverter, motor, wiring harness agents", "k vehicle-equivalent components"],
            ["3", "OEM assembly", "BYD, other Chinese, UK, US, European, Korean, Japanese OEM groups", "k vehicles assembled"],
            ["4", "Markets", "China, Europe, UK, USA, Japan, ROW demand agents", "Weekly EV demand in GWh and k vehicles"],
        ],
        [1.0, 3.0, 7.0, 5.0],
    )
    caption(doc, "Table 1. Model layers and principal outputs.")

    heading(doc, "3. ODD Protocol", 1)
    heading(doc, "3.1 Overview", 2)
    para(
        doc,
        "Following the ODD protocol, the model purpose is to explain dynamic supply-chain vulnerability under critical mineral, cell, Tier-1, and OEM disruption scenarios. "
        "The main entities are ABM agents and SD stocks. The temporal resolution is one week. State variables include inventories, capacity, demand, prices, backlog, chemistry share, and shock multipliers."
    )
    heading(doc, "3.2 Design Concepts", 2)
    bullet(doc, "Emergence: aggregate production loss emerges from local production, ordering, inventory, and shock-recovery decisions.")
    bullet(doc, "Adaptation: agents adapt output to price premiums, shortages, financial profiles, and archetype-specific behavioural rules.")
    bullet(doc, "Sensing: agents sense SD availability fractions, component/mineral prices, downstream demand, and their own inventory or backlog.")
    bullet(doc, "Stochasticity: mineral flows include weekly supply volatility; market demand contains noise around the growth trend.")
    bullet(doc, "Interaction: agents interact indirectly through shared SD stocks and directly through downstream demand allocation.")
    heading(doc, "3.3 Details", 2)
    para(
        doc,
        "Initialisation uses 2023 baseline values from public data sources. The schedule is fixed: shocks, SD signal computation, agent stepping by tier, market demand update, SD flow aggregation, SD update, and metric recording. "
        "Submodels include shock handling, Leontief production constraints, price formation, LFP substitution, capacity investment, order-up-to logic, market demand response, and policy effects."
    )

    heading(doc, "4. Mathematical Specification", 1)
    heading(doc, "4.1 Time Step and Stock-Flow Updating", 2)
    para(doc, "The SD layer uses discrete-time Euler integration with dt = 1 week. For a generic stock S:")
    eq(doc, "S_{t+1} = max(0, S_t + In_t - Out_t)")
    para(doc, "Mineral and component inventories are expressed mainly as weeks of baseline supply, which makes safety-stock targets and shortfall thresholds directly interpretable.")

    heading(doc, "4.2 Transport and Perception Delays", 2)
    para(doc, "Mineral supply enters an in-transit FIFO pipeline before reaching usable inventory:")
    eq(doc, "Arrival_{m,t} = Pipeline_m.pop_left();  Pipeline_m.append(Supply_{m,t})")
    para(doc, "Agents do not observe true inventory directly. A first-order measurement lag creates measured stock:")
    eq(doc, "MeasuredStock_{m,t+1} = MeasuredStock_{m,t} + lambda_m (Stock_{m,t} - MeasuredStock_{m,t})")

    heading(doc, "4.3 Availability and Leontief Production", 2)
    para(doc, "Availability for input i is computed as the measured stock relative to target weeks of supply:")
    eq(doc, "a_{i,t} = min(1, MeasuredStock_{i,t} / TargetStock_i)")
    para(doc, "Downstream production follows a Leontief-style bottleneck condition. For an agent with relevant input set I:")
    eq(doc, "Constraint_t = min_{i in I}(a_{i,t}^{d_i})")
    eq(doc, "Output_t = min(Capacity_t, Demand_t + BacklogClearance_t, Inventory_t) x ShockMultiplier_t x Constraint_t")

    heading(doc, "4.4 Shock Dynamics", 2)
    para(doc, "A shock is represented by target, start week, end week, and severity. During the active interval:")
    eq(doc, "ShockMultiplier_{j,t} = max(0, 1 - Severity_j x (1 - PolicyMitigation_j))")
    para(doc, "When a shock resolves, agents recover gradually according to their recovery-rate parameter and archetype rules.")

    heading(doc, "4.5 Price Formation", 2)
    para(doc, "Commodity prices respond to scarcity using a bounded softplus function. The target price for mineral m is:")
    eq(doc, "p*_{m,t} = 1 + alpha x softplus_beta(1 - a_{m,t})")
    eq(doc, "softplus_beta(x) = ln(1 + exp(beta x)) / beta")
    eq(doc, "p_{m,t+1} = clip(p_{m,t} + gamma_m (p*_{m,t} - p_{m,t}), p_min, p_max)")
    para(doc, "This avoids singular behaviour when stocks approach zero while preserving a smooth scarcity response.")

    heading(doc, "4.6 Logistic Chemistry Substitution", 2)
    para(doc, "Cobalt price pressure shifts the target LFP share through a logistic function of delayed cobalt price:")
    eq(doc, "pC_delayed_{t+1} = pC_delayed_t + lambda_C (pC_t - pC_delayed_t)")
    eq(doc, "LFP_Target_t = LFP_min + (LFP_max - LFP_min) / (1 + exp[-beta(log(pC_delayed_t) - log(p_mid))])")
    eq(doc, "LFPShare_{t+1} = LFPShare_t + k_LFP (LFP_Target_t - LFPShare_t)")

    heading(doc, "4.7 Market Demand", 2)
    para(doc, "Each market has an exogenous growth trend modified by price and availability effects:")
    eq(doc, "TrendDemand_{r,t+1} = TrendDemand_{r,t}(1 + g_r)")
    eq(doc, "Demand_{r,t} = TrendDemand_{r,t} x clip(1 + epsilon_r(PriceSignal_t - 1), 0.50, 1.20) x AvailabilityEffect_t")

    heading(doc, "4.8 Bullwhip Index", 2)
    para(doc, "Tier-1 order rates are compared with realised demand to produce a smoothed bullwhip indicator:")
    eq(doc, "RawBullwhip_t = OrderRate_t / max(Demand_t, epsilon)")
    eq(doc, "Bullwhip_t = rho Bullwhip_{t-1} + (1-rho) RawBullwhip_t")

    heading(doc, "5. Logistics and Operational Logic", 1)
    para(
        doc,
        "The operational logic is deliberately supply-chain oriented. Materials are delayed by transport pipelines; component agents hold inventory and place orders with lead times; OEMs assemble vehicles only when packs, inverters, motors, and harnesses are jointly available; market agents adjust realised demand to price and availability."
    )
    table(
        doc,
        ["Operational Element", "Implementation Logic"],
        [
            ["Transport", "Mineral inflows are placed in FIFO pipelines whose lengths approximate supply-chain lead times."],
            ["Inventory", "Stocks are held in weeks of baseline consumption for minerals and vehicle-equivalent units for components."],
            ["Ordering", "Tier-1 suppliers use order-up-to rules with bullwhip amplification and archetype-specific modifications."],
            ["Recovery", "Shocked agents recover gradually after shock end; financial resilience modifies absorption and recovery."],
            ["Policy", "Policy packages operate on both ABM agents and SD parameters, including buffers, recovery, mitigation, recycling, demand growth, and price smoothing."],
        ],
        [4.0, 12.0],
    )

    heading(doc, "6. SD-ABM Communication", 1)
    para(doc, "The hybrid model communicates in a fixed weekly order. This avoids circular simultaneity inside a single time step while allowing feedback across weeks.")
    table(
        doc,
        ["Step", "Direction", "Information Exchanged"],
        [
            ["1", "Scenario -> Agents", "Scheduled shocks change target agents' shock multipliers."],
            ["2", "SD -> ABM", "SD computes input_fractions and price_signals from measured stocks and prices."],
            ["3", "ABM internal", "Mineral, cell, Tier-1, OEM, and market agents step in sequence."],
            ["4", "ABM -> SD", "Hybrid model aggregates outputs into mineral inflows/outflows, component inflows/outflows, cell output, demand, backlog, and order rates."],
            ["5", "SD internal", "SD updates transport pipelines, stocks, prices, LFP share, capacity, demand/backlog, and bullwhip."],
            ["6", "Model -> Results", "Weekly metrics are recorded for dashboard, validation, and policy evaluation."],
        ],
        [1.0, 3.0, 12.0],
    )
    caption(doc, "Table 2. Weekly SD-ABM communication schedule.")

    heading(doc, "7. Agent Classes and Behavioural Archetypes", 1)
    table(
        doc,
        ["Tier", "Base Agent", "Archetypes", "Behavioural Meaning"],
        [
            ["Minerals", "MineralSupplierAgent", "StateBacked, WesternMiner, GreenfieldBuilder", "Different price response, output floors, strategic restrictions, mothballing, and debt-driven utilisation."],
            ["Cells", "CellManufacturerAgent", "PlatformLeader, HyperScaleChallenger, IncumbentUnderPressure", "Different demand-pull versus push strategies, LFP adaptation, liquidity sensitivity, and growth rates."],
            ["Tier-1", "Tier1SupplierAgent", "PremiumPowerElectronics, EstablishedVolumeSupplier, BatteryPackIntegrator", "Different order timing, price-sensitive procurement, JIT pass-through, and buffer management."],
            ["OEM", "OEMAgent", "ProfitableEstablishedOEM, TransitioningLegacyOEM, EVNativeScaleAspirant, PrecommercialStartup", "Different margin response, backlog clearance, fallback to ICE capacity, demand chasing, and cash constraints."],
            ["Market", "MarketAgent", "Regional market settings", "Demand growth, price elasticity, backlog sensitivity, availability floor, and average kWh per vehicle."],
        ],
        [2.0, 3.0, 5.0, 6.0],
    )

    heading(doc, "8. Causal-Loop Diagram", 1)
    para(
        doc,
        "The causal-loop structure can be read as a set of coupled balancing and reinforcing mechanisms. "
        "The table below serves as a text causal-loop diagram suitable for inclusion in the Word paper; it corresponds to the model's SD feedback loops and ABM ordering behaviour."
    )
    causal_loop_table(doc)

    heading(doc, "9. Assumptions", 1)
    bullet(doc, "Weekly time is sufficient to capture industrial supply-chain delays without modelling daily logistics.")
    bullet(doc, "Stocks expressed in weeks of baseline supply are comparable across materials and components.")
    bullet(doc, "Short-run production is bottlenecked by the scarcest required input, represented through a Leontief-style constraint.")
    bullet(doc, "Financial and behavioural archetypes can approximate firm heterogeneity without full balance-sheet simulation.")
    bullet(doc, "Market demand growth is exogenous but modified by endogenous price and availability signals.")
    bullet(doc, "Policy interventions affect buffers, mitigation, recovery, recycling/offtake, demand growth, and price smoothing, not macroeconomic conditions.")
    bullet(doc, "The UK OEM group represents JLR, MINI Oxford, Vauxhall Ellesmere Port, and related UK EV production exposure as one aggregate endpoint.")

    heading(doc, "10. Data Sources and Parameterisation", 1)
    para(
        doc,
        "Parameters are set from public sources where available, then translated into weekly simulation units. When direct weekly data do not exist, values are selected from industry rules of thumb, documented lead times, or conservative calibration choices. The model favours transparent, inspectable parameters over automated black-box fitting."
    )
    table(
        doc,
        ["Parameter Area", "Main Source Logic", "Examples"],
        [
            ["Critical minerals", "USGS Mineral Commodity Summaries and IEA Critical Minerals Market Review.", "Cobalt DRC share 70%; graphite China share 79%; REE China processing share 85%; mineral safety stocks 4-12 weeks."],
            ["Cell makers", "IEA Global EV Outlook, BNEF, SNE Research, company reports.", "Global deployed cells 822 GWh in 2023; CATL 37%; BYD cells 14%; AESC UK small 2023 capacity."],
            ["Tier-1 components", "Automotive lead-time evidence, company disclosures, technology reports.", "Harness lead time 6 weeks; inverter lead time 16 weeks; motor lead time 12 weeks; battery pack lead time 4 weeks."],
            ["OEM groups", "IEA EV production totals, SMMT UK production data, company annual reports.", "UK OEM 175k vehicles/year; BYD 1.575m BEVs/year; other Chinese OEM residual volume."],
            ["Market demand", "IEA EV demand and regional registration data.", "Regional GWh demand, average kWh per vehicle, annual growth rates, price elasticities."],
            ["Prices and delays", "Benchmark/BNEF/industry reports and calibration to plausible response speeds.", "Commodity price adjustment 0.05/week; LFP shift speed 0.003/week; cell build time 104 weeks."],
            ["Validation data", "Generated validation artifacts and real ONS/SMMT time-series alignment where available.", "Validation checks, scenario metrics, monthly UK car production index comparison."],
        ],
        [3.0, 7.0, 6.0],
    )
    caption(doc, "Table 3. Parameterisation sources and examples.")

    heading(doc, "11. Shock and Policy Scenario Logic", 1)
    para(
        doc,
        "Shock scenarios are dictionaries containing target, start_week, end_week, and severity. They can target mineral suppliers, cell makers, Tier-1 suppliers, or OEMs. The dashboard now also allows users to construct a custom shock by selecting shock type, target, timing, duration, and severity."
    )
    table(
        doc,
        ["Scenario Type", "Example Targets", "Interpretation"],
        [
            ["Mineral supply", "cobalt_drc, graphite_chn, ree_chn, sic_wolfspeed", "Political disruption, export controls, mine outage, or processing bottleneck."],
            ["Cell manufacturing", "cell_catl, cell_byd_cells, cell_lg_es, cell_aesc_uk", "Plant disruption, power rationing, supplier concentration risk, or regional production shock."],
            ["Tier-1 components", "t1_harness, t1_inverter, t1_motor, t1_battery_pack", "Harness logistics disruption, SiC inverter bottleneck, motor magnet shortage, pack assembly constraint."],
            ["OEM assembly", "oem_uk_oem and other OEM groups", "Border friction, labour/plant disruption, throughput shock, or policy compliance friction."],
        ],
        [3.0, 5.0, 8.0],
    )

    heading(doc, "12. Validation Approach", 1)
    para(
        doc,
        "Validation combines structural checks, face-validity scenario checks, behavioural/archetype tests, numerical sanity checks, and comparison with real time-series where available. "
        "The aim is not to fit every historical month exactly, but to ensure that mechanisms produce directionally credible, bounded, and explainable behaviour."
    )
    bullet(doc, "Structural validation: all scenario targets exist, severities are bounded in [0,1], and start/end windows are valid.")
    bullet(doc, "Behavioural validation: archetype rules activate under the expected price, inventory, and shock states.")
    bullet(doc, "Scenario validation: DRC cobalt, Ukraine harness, China graphite/REE, CATL disruption, SiC bottleneck, compound shock, and policy variants are checked against expected propagation pathways.")
    bullet(doc, "Numerical validation: stocks, prices, demand, backlog, and production remain finite and bounded.")
    bullet(doc, "Historical alignment: UK OEM model output can be indexed against ONS/SMMT monthly UK car production where validation data are available.")

    heading(doc, "13. Limitations", 1)
    bullet(doc, "Aggregation: each tier is represented by aggregate agents; plant-level and contract-level flows are not fully resolved.")
    bullet(doc, "Cell routing: the model contains OEM cell-source metadata, but some flows still operate through aggregate cell market-share demand rather than a full bilateral OEM-cell network.")
    bullet(doc, "Uncertainty: single-run scenarios include stochastic disturbances, but full Monte Carlo confidence intervals are not generated by default in the app.")
    bullet(doc, "Policy realism: policy packages are simplified mechanisms, not full fiscal, planning, grid, workforce, or trade models.")
    bullet(doc, "Data gaps: public data often mix deliveries, registrations, production, nameplate capacity, and deployed GWh; parameters are therefore harmonised with transparent assumptions.")
    bullet(doc, "Demand: regional adoption trends are simplified and cannot represent all macroeconomic, consumer, charging infrastructure, or regulatory feedbacks.")
    bullet(doc, "Substitution: LFP substitution captures cobalt avoidance, but other chemistry shifts, recycling chemistry quality, and mineral-by-mineral technical constraints remain simplified.")

    heading(doc, "14. Reproducibility", 1)
    para(
        doc,
        "The model is implemented in Python. Core implementation files are model/sd_model.py, model/agents.py, model/hybrid_model.py, model/shocks.py, model/policies.py, run_simulation.py, validate_model.py, and streamlit_app.py. "
        "The app can be launched with `python -m streamlit run streamlit_app.py`. Validation artifacts are stored in the results directory."
    )

    heading(doc, "References", 1)
    refs = [
        "AESC (2023). Envision AESC Sunderland gigafactory announcements and public capacity statements.",
        "BloombergNEF (2023). Battery Price Survey and battery-market outlook data.",
        "Bonabeau, E. (2002). Agent-based modeling: Methods and techniques for simulating human systems. PNAS, 99(suppl. 3), 7280-7287.",
        "Forrester, J. W. (1961). Industrial Dynamics. MIT Press.",
        "IEA (2023). Critical Minerals Market Review. International Energy Agency.",
        "IEA (2024). Global EV Outlook. International Energy Agency.",
        "Lee, H. L., Padmanabhan, V., and Whang, S. (1997). Information distortion in a supply chain: The bullwhip effect. Management Science, 43(4), 546-558.",
        "Leoni AG (2022). Annual report and disclosures on Ukraine harness disruption.",
        "SMMT (2024). UK vehicle production and registration data.",
        "Sterman, J. D. (2000). Business Dynamics: Systems Thinking and Modeling for a Complex World. McGraw-Hill.",
        "USGS (2024). Mineral Commodity Summaries. United States Geological Survey.",
        "Yole Intelligence (2023). Power electronics and SiC market reports.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.first_line_indent = Cm(-0.7)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(ref)
        r.font.size = Pt(9.5)
        r.font.color.rgb = TEXT

    doc.save(OUT_PATH)
    print(f"Saved {OUT_PATH}")
    print(f"Size {os.path.getsize(OUT_PATH) // 1024} KB")


if __name__ == "__main__":
    build_document()
