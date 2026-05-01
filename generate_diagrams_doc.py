"""
Generate CLD + ODD Word document.
Produces:  results/EV_Supply_Chain_CLD_ODD.docx
"""

import io
import os
import math
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyArrowPatch
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── helpers ───────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_colour):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tcPr.append(shd)


def _heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run()
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    return p


def _body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    return p


def _bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.6)
    p.paragraph_format.space_after = Pt(2)
    return p


def _table_header_row(table, headers, bg="1E3A5F"):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(cell, bg)


def _add_row(table, values, bold_first=False, bg=None):
    row = table.add_row()
    for i, v in enumerate(values):
        cell = row.cells[i]
        cell.text = str(v)
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        if bold_first and i == 0:
            cell.paragraphs[0].runs[0].font.bold = True
        if bg:
            _set_cell_bg(cell, bg)
    return row


def _fig_to_docx(doc, fig, width_inches=6.2, caption=None):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=180, bbox_inches="tight")
    buf.seek(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Inches(width_inches))
    plt.close(fig)
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].font.size = Pt(9)
        cp.runs[0].font.italic = True
        cp.runs[0].font.color.rgb = RGBColor(0x47, 0x55, 0x69)


# ══════════════════════════════════════════════════════════════════════════════
# FIGURE 1 — Causal Loop Diagram
# ══════════════════════════════════════════════════════════════════════════════

def _arrow(ax, x0, y0, x1, y1, label, polarity, color, lw=1.4, rad=0.0,
           label_offset=(0, 0)):
    """Draw a curved arrow with polarity label."""
    style = f"Arc3,rad={rad}"
    ax.annotate(
        "",
        xy=(x1, y1), xytext=(x0, y0),
        arrowprops=dict(
            arrowstyle="-|>",
            color=color,
            lw=lw,
            connectionstyle=style,
        ),
    )
    mx = (x0 + x1) / 2 + label_offset[0]
    my = (y0 + y1) / 2 + label_offset[1]
    sign_color = "#16a34a" if polarity == "+" else "#dc2626"
    ax.text(mx, my, polarity, ha="center", va="center",
            fontsize=10, fontweight="bold", color=sign_color,
            bbox=dict(boxstyle="round,pad=0.15", fc="white", ec="none", alpha=0.85))
    if label:
        ax.text(mx + 0.02, my - 0.14, label, ha="center", va="top",
                fontsize=6.5, color="#475569", style="italic")


def build_cld() -> plt.Figure:
    """
    Nodes (x, y) on a 10×8 canvas.

    Stocks / variables:
      Mineral Supply  →  Mineral Stock  →  Mineral Price  →  LFP Share
                                        ↘  Price Signal   →  OEM Demand
      Mineral Stock   →  Cell Output    →  Cell Inventory →  OEM Production
      OEM Production  →  Backlog        →  OEM Demand (F4)
      Cell Capacity   →  Cell Output    (investment loop F3)
      Utilisation     →  Capex          →  Cell Capacity  (F3)
    """
    fig, ax = plt.subplots(figsize=(13, 9))
    ax.set_xlim(0, 13)
    ax.set_ylim(0, 9)
    ax.axis("off")
    ax.set_facecolor("#f8fafc")
    fig.patch.set_facecolor("#f8fafc")

    # ── Node positions ──────────────────────────────────────────────────────
    nodes = {
        # Row 0 — minerals
        "Mineral\nSupply":      (1.2, 7.5),
        "Mineral\nStock":       (3.5, 7.5),
        "Mineral\nPrice":       (6.0, 7.5),
        "LFP\nShare":           (8.5, 7.5),
        "Cobalt\nSensitivity":  (8.5, 6.2),

        # Row 1 — cell tier
        "Cell\nOutput":         (3.5, 5.8),
        "Cell\nInventory":      (6.0, 5.8),
        "Cell\nCapacity":       (1.2, 5.8),
        "Utilisation":          (1.2, 4.4),
        "Capex":                (1.2, 3.0),

        # Row 2 — OEM / demand
        "Price\nSignal":        (6.0, 4.4),
        "OEM\nProduction":      (8.8, 4.4),
        "OEM\nBacklog":         (11.2, 4.4),
        "EV\nDemand":           (11.2, 6.0),
        "Market\nAvailability": (9.8, 6.0),
    }

    # ── Draw nodes ──────────────────────────────────────────────────────────
    STOCK_NODES = {"Mineral\nStock", "Cell\nInventory", "Cell\nCapacity",
                   "OEM\nBacklog", "EV\nDemand"}
    for name, (x, y) in nodes.items():
        is_stock = name in STOCK_NODES
        fc = "#dbeafe" if is_stock else "#f0fdf4"
        ec = "#2563eb" if is_stock else "#16a34a"
        box = mpatches.FancyBboxPatch(
            (x - 0.68, y - 0.38), 1.36, 0.76,
            boxstyle="round,pad=0.08",
            linewidth=1.6 if is_stock else 1.2,
            edgecolor=ec, facecolor=fc,
        )
        ax.add_patch(box)
        ax.text(x, y, name, ha="center", va="center",
                fontsize=7.5, fontweight="bold" if is_stock else "normal",
                color="#0f172a", multialignment="center")

    # ── Causal links ────────────────────────────────────────────────────────
    # F1 – Supply-demand price loop (Balancing)
    _arrow(ax, 1.88, 7.5,  2.82, 7.5,  "",  "+", "#2563eb", rad=0.0, label_offset=(0, 0.18))
    _arrow(ax, 4.18, 7.5,  5.32, 7.5,  "",  "+", "#2563eb", rad=0.0, label_offset=(0, 0.18))
    # Mineral Price → back to Mineral Supply (B loop closing)
    _arrow(ax, 6.0, 7.12, 3.5, 7.12,  "F1",  "+", "#2563eb", rad=-0.18,
           label_offset=(0.1, -0.28))

    # F2 – Chemistry substitution (Balancing)
    _arrow(ax, 6.68, 7.5,  7.82, 7.5,  "",  "+", "#7c3aed", rad=0.0, label_offset=(0, 0.18))
    _arrow(ax, 8.5, 7.12,  8.5, 6.58,  "",  "+", "#7c3aed", rad=0.0, label_offset=(0.25, 0))
    _arrow(ax, 8.5, 5.82,  6.0, 7.12,  "F2",  "-", "#7c3aed", rad=0.12,
           label_offset=(0.0, 0.2))

    # Mineral Stock → Cell Output
    _arrow(ax, 3.5, 7.12, 3.5, 6.18, "", "+", "#0369a1", rad=0.0, label_offset=(0.28, 0))
    # LFP / Cobalt sensitivity → Cell Output
    _arrow(ax, 8.5, 5.82, 3.5, 5.82, "", "-", "#7c3aed", rad=0.08,
           label_offset=(0, -0.2))

    # Cell Output → Cell Inventory
    _arrow(ax, 4.18, 5.8, 5.32, 5.8, "", "+", "#0369a1", rad=0.0, label_offset=(0, 0.2))
    # Cell Inventory → OEM Production (via availability)
    _arrow(ax, 6.68, 5.8, 8.12, 4.62, "", "+", "#0369a1", rad=0.12,
           label_offset=(0.3, 0.1))

    # Cell Output → Utilisation
    _arrow(ax, 2.82, 5.8, 1.88, 5.8, "", "+", "#0369a1", rad=0.0, label_offset=(0, 0.2))
    # Cell Capacity → Cell Output
    _arrow(ax, 1.88, 5.8, 2.82, 5.8, "", "+", "#15803d", rad=-0.28,
           label_offset=(0, -0.28))
    # Utilisation → Capex  (F3 trigger)
    _arrow(ax, 1.2, 4.02, 1.2, 3.38, "", "+", "#15803d", rad=0.0, label_offset=(0.28, 0))
    # Capex → Cell Capacity
    _arrow(ax, 1.2, 2.62, 1.2, 5.42, "F3", "+", "#15803d", rad=0.38,
           label_offset=(-0.55, 0))

    # Mineral Price → Price Signal
    _arrow(ax, 6.0, 7.12, 6.0, 4.78, "", "+", "#dc2626", rad=0.0, label_offset=(0.28, 0))
    # Price Signal → OEM Production
    _arrow(ax, 6.68, 4.4, 8.12, 4.4, "", "-", "#dc2626", rad=0.0, label_offset=(0, 0.2))
    # Price Signal → EV Demand
    _arrow(ax, 6.0, 5.42, 6.0, 5.42, "", "", "#dc2626")  # placeholder – handled below

    # OEM Production → OEM Backlog
    _arrow(ax, 9.48, 4.4, 10.52, 4.4, "", "+", "#ea580c", rad=0.0, label_offset=(0, 0.2))
    # OEM Backlog → EV Demand (F4 Balancing)
    _arrow(ax, 11.2, 4.78, 11.2, 5.62, "", "+", "#ea580c", rad=0.0, label_offset=(0.28, 0))
    # EV Demand → OEM Production
    _arrow(ax, 10.52, 6.0, 9.48, 4.62, "F4", "+", "#ea580c", rad=0.08,
           label_offset=(0.35, 0.0))
    # Market Availability → OEM Production
    _arrow(ax, 9.8, 5.62, 9.1, 4.78, "", "+", "#ea580c", rad=0.0,
           label_offset=(0.32, 0))
    # Price Signal → Market Availability
    _arrow(ax, 6.68, 4.4, 9.12, 5.82, "", "-", "#dc2626", rad=-0.12,
           label_offset=(0.2, 0.2))

    # F5 Bullwhip — Tier-1 order amplification
    _arrow(ax, 3.5, 5.42, 3.5, 4.58, "F5", "+", "#b45309", rad=0.28,
           label_offset=(-0.38, 0))
    _arrow(ax, 3.5, 4.58, 3.5, 5.42, "", "-", "#b45309", rad=0.28,
           label_offset=(0.38, 0))

    # ── Legend ──────────────────────────────────────────────────────────────
    legend_x, legend_y = 10.0, 2.6
    ax.text(legend_x, legend_y + 0.5, "Legend", fontsize=8, fontweight="bold",
            color="#0f172a")
    for i, (label, fc, ec) in enumerate([
        ("Stock / Accumulator", "#dbeafe", "#2563eb"),
        ("Auxiliary / Rate",    "#f0fdf4", "#16a34a"),
    ]):
        bx = mpatches.FancyBboxPatch(
            (legend_x, legend_y - i * 0.5), 0.7, 0.38,
            boxstyle="round,pad=0.06", lw=1.2, ec=ec, fc=fc)
        ax.add_patch(bx)
        ax.text(legend_x + 0.85, legend_y + 0.19 - i * 0.5, label,
                fontsize=7.5, va="center", color="#0f172a")

    loop_colors = [
        ("#2563eb", "F1  Supply–demand price  (B)"),
        ("#7c3aed", "F2  Chemistry substitution (B)"),
        ("#15803d", "F3  Cell capacity investment (R→B)"),
        ("#ea580c", "F4  Demand–adoption (B)"),
        ("#b45309", "F5  Bullwhip amplification"),
    ]
    for i, (col, lbl) in enumerate(loop_colors):
        ax.plot([legend_x, legend_x + 0.7],
                [legend_y - 1.2 - i * 0.42, legend_y - 1.2 - i * 0.42],
                color=col, lw=2)
        ax.text(legend_x + 0.85, legend_y - 1.2 - i * 0.42,
                lbl, fontsize=7.5, va="center", color="#0f172a")

    ax.text(legend_x, legend_y - 3.5,
            "+ = same direction\n− = opposite direction\nB = Balancing loop\nR = Reinforcing loop",
            fontsize=7.5, color="#475569", va="top", linespacing=1.5)

    ax.set_title(
        "Figure 1 — Causal Loop Diagram: EV Supply Chain System Dynamics Layer",
        fontsize=11, fontweight="bold", color="#0f172a", pad=12,
    )
    fig.tight_layout(pad=0.5)
    return fig


# ══════════════════════════════════════════════════════════════════════════════
# FIGURE 2 — ABM Tier Architecture diagram
# ══════════════════════════════════════════════════════════════════════════════

def build_abm_arch() -> plt.Figure:
    """Four-tier ABM agent architecture with archetype breakdown."""
    fig, ax = plt.subplots(figsize=(13, 7.5))
    ax.set_xlim(0, 13)
    ax.set_ylim(0, 7.5)
    ax.axis("off")
    ax.set_facecolor("#f8fafc")
    fig.patch.set_facecolor("#f8fafc")

    TIER_META = [
        (0, "Tier 0 — Mineral Suppliers",  "#fef3c7", "#d97706",
         [("StateBacked\n(4 agents)",   "graphite_chn, ree_chn,\ncobalt_other, sic_china",
           "Restricts when price > 1.8×"),
          ("WesternMiner\n(5 agents)",  "lithium_aus/chl, cobalt_drc,\nsic_coherent, sic_other",
           "Mothballs after 12 wk below 0.85×"),
          ("GreenfieldBuilder\n(3 agents)", "lithium_other, ree_other,\nsic_wolfspeed",
           "Permanent degradation after distress")]),

        (1, "Tier 1 — Cell Manufacturers", "#dbeafe", "#2563eb",
         [("PlatformLeader\n(2 agents)",   "catl (37 %), byd_cells (14 %)",
           "Full-capacity stockpile; fast LFP shift"),
          ("HyperScaleChallenger\n(2 agents)", "calb (5 %), others_cells (12.8 %)",
           "Always max output; growth penalty if shortfall"),
          ("IncumbentUnderPressure\n(4 agents)", "lg_es, panasonic, samsung_sdi, sk_on",
           "Demand-pull; slow LFP shift; share erosion")]),

        (2, "Tier 2 — Sub-system Suppliers", "#f3e8ff", "#7c3aed",
         [("BatteryPackIntegrator\n(1 agent)",  "battery_pack",
           "Pure JIT — zero amplification"),
          ("PremiumPowerElectronics\n(1 agent)", "inverter (16-wk lead)",
           "Forward demand projection; SiC price deferral"),
          ("EstablishedVolumeSupplier\n(2 agents)", "motor, harness",
           "Production smoothing; standard bullwhip")]),

        (3, "Tier 3 — OEMs",               "#dcfce7", "#15803d",
         [("TransitioningLegacyOEM\n(3 agents)", "german_oem, us_oem, uk_oem",
           "ICE fallback when margin < 85 %"),
          ("EVNativeScaleAspirant\n(1 agent)",   "other_chinese_oem",
           "Demand-elastic; proximity boost"),
          ("ProfitableEstablishedOEM\n(2 agents)", "korean_oem, japanese_oem",
           "Buffer-first; full prod when inv > 50 % target")]),
    ]

    tier_ys = [6.5, 4.7, 2.9, 1.1]
    tier_h  = 1.4

    for tier_idx, title, fc, ec, archetypes in TIER_META:
        y0 = tier_ys[tier_idx] - tier_h / 2
        # Tier background band
        band = mpatches.FancyBboxPatch(
            (0.15, y0), 12.7, tier_h,
            boxstyle="round,pad=0.08", lw=1.5, ec=ec, fc=fc, alpha=0.45)
        ax.add_patch(band)
        ax.text(0.35, tier_ys[tier_idx], title,
                fontsize=9, fontweight="bold", va="center", color="#0f172a")

        # Archetype boxes
        box_w = 3.7
        for a_idx, (arch_name, agents, behaviour) in enumerate(archetypes):
            bx = 3.8 + a_idx * (box_w + 0.2)
            by = tier_ys[tier_idx]
            rect = mpatches.FancyBboxPatch(
                (bx, by - 0.52), box_w, 1.04,
                boxstyle="round,pad=0.07", lw=1.2, ec=ec,
                fc="white", alpha=0.9)
            ax.add_patch(rect)
            ax.text(bx + box_w / 2, by + 0.26, arch_name,
                    ha="center", va="center", fontsize=7.2,
                    fontweight="bold", color=ec)
            ax.text(bx + box_w / 2, by - 0.04, agents,
                    ha="center", va="center", fontsize=6.2,
                    color="#334155", style="italic")
            ax.text(bx + box_w / 2, by - 0.34, behaviour,
                    ha="center", va="center", fontsize=6.0,
                    color="#64748b")

    # Vertical flow arrows between tiers
    for yi in range(len(tier_ys) - 1):
        y_from = tier_ys[yi] - tier_h / 2 - 0.04
        y_to   = tier_ys[yi + 1] + tier_h / 2 + 0.04
        ax.annotate(
            "", xy=(2.1, y_to), xytext=(2.1, y_from),
            arrowprops=dict(arrowstyle="-|>", color="#64748b", lw=1.4),
        )
        ax.text(2.3, (y_from + y_to) / 2, "supply\nflow",
                fontsize=6.5, color="#64748b", va="center", style="italic")

    # SD coupling annotation
    ax.annotate(
        "", xy=(12.5, 4.7), xytext=(12.5, 1.1),
        arrowprops=dict(arrowstyle="<->", color="#0369a1", lw=1.6,
                        linestyle="dashed"),
    )
    ax.text(12.6, 2.9, "SD\ncoupling", fontsize=7, color="#0369a1",
            va="center", style="italic")

    ax.set_title(
        "Figure 2 — ABM Four-Tier Agent Architecture with Archetypes",
        fontsize=11, fontweight="bold", color="#0f172a", pad=10,
    )
    fig.tight_layout(pad=0.4)
    return fig


# ══════════════════════════════════════════════════════════════════════════════
# BUILD WORD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

def build_doc():
    doc = Document()

    # Page margins
    for sect in doc.sections:
        sect.top_margin    = Cm(2.0)
        sect.bottom_margin = Cm(2.0)
        sect.left_margin   = Cm(2.5)
        sect.right_margin  = Cm(2.5)

    # ── Title ─────────────────────────────────────────────────────────────────
    title = doc.add_heading(
        "EV Supply Chain Hybrid Model: Causal Loop Diagram & ODD Protocol", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    sub = doc.add_paragraph(
        "System Dynamics Causal Loop Diagram  ·  Agent-Based Model ODD Protocol")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(10)
    sub.runs[0].font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # PART A — CAUSAL LOOP DIAGRAM
    # ══════════════════════════════════════════════════════════════════════════
    _heading(doc, "Part A — System Dynamics Causal Loop Diagram", 1)
    _body(doc,
        "The System Dynamics (SD) layer of the hybrid model governs six stocks of "
        "critical minerals, cell inventory, cell capacity, component stocks, chemistry mix, "
        "commodity prices, and aggregate EV demand. The causal loop diagram (CLD) below "
        "maps the five principal feedback structures that drive system behaviour.")
    doc.add_paragraph()

    fig_cld = build_cld()
    _fig_to_docx(doc, fig_cld, width_inches=6.2,
                 caption="Figure 1. Causal Loop Diagram — EV Supply Chain System Dynamics Layer. "
                         "Blue boxes = stocks (accumulators); green boxes = auxiliary variables. "
                         "B = Balancing loop; R = Reinforcing loop.")

    doc.add_paragraph()
    _heading(doc, "A.1  Feedback Loop Descriptions", 2)

    LOOPS = [
        ("F1", "Supply–Demand Price Loop (Balancing)",
         "#2563eb",
         "Mineral Supply → (+) Mineral Stock → (+) Mineral Price → (+) Mineral Supply. "
         "When mineral stocks fall below target (e.g., due to a DRC cobalt shock), the "
         "measured stock deficit drives price upward through a softplus scarcity function. "
         "Higher prices stimulate supply recovery, closing the gap. The measurement lag "
         "(τ ≈ 2 weeks) and transport delay (4–10 weeks) introduce a phase shift that can "
         "produce transient price overshoots before the loop stabilises."),

        ("F2", "Chemistry Substitution Loop (Balancing)",
         "#7c3aed",
         "Mineral Price (cobalt) → (+) LFP Share → (−) Cobalt Sensitivity → (−) Mineral Price. "
         "When the cobalt price index rises above 1.3×, cell manufacturers shift toward LFP "
         "chemistry (which is cobalt-free). The LFP target is set by a logistic function of "
         "log-cobalt-price. A 4-week first-order lag on cobalt price perception models "
         "managerial re-tooling delay. PlatformLeader agents shift at 0.5 %/week; "
         "IncumbentUnderPressure agents shift at 0.1 %/week, producing archetype-driven "
         "heterogeneity in this loop's speed."),

        ("F3", "Cell Capacity Investment Cycle (Reinforcing → Balancing)",
         "#15803d",
         "Cell Output → (+) Utilisation → (+) Capex → (+) Cell Capacity → (+) Cell Output. "
         "When utilisation exceeds 85 %, investment is triggered. A 26-week planning queue "
         "feeds a 3-stage Erlang construction pipeline (mean build = 104 weeks). This "
         "produces a realistic 2–3 year capital cycle. Investment is capped at 30 %/yr to "
         "prevent unrealistic instantaneous expansion. The loop is initially reinforcing "
         "(more capacity → more output → more utilisation until saturation), then balancing "
         "once capacity overtakes demand growth."),

        ("F4", "Demand–Adoption Feedback Loop (Balancing)",
         "#ea580c",
         "OEM Production → (+) OEM Backlog → (+) EV Demand → (+) OEM Production. "
         "Rising backlogs signal unmet demand, pulling more vehicles into production. "
         "The market agent converts backlog-adjusted demand into a weekly OEM production "
         "target. Price elasticity (−0.36) and an availability floor (0.46) moderate "
         "the strength of this loop. Under Brexit friction or CATL disruption, the backlog "
         "accumulates faster than OEM capacity can absorb, stretching recovery timelines."),

        ("F5", "Bullwhip Amplification (Tracking structure)",
         "#b45309",
         "Tier-1 order variability is tracked via an EWMA of the ratio of order variance "
         "to demand variance (bullwhip index). This is not a closed feedback loop but a "
         "structural tendency: the order-up-to inventory policy with a bullwhip factor "
         "(default 1.25×) amplifies demand signals up the supply chain. The "
         "BatteryPackIntegrator archetype bypasses this (exact JIT ordering), while "
         "PremiumPowerElectronics 16-week forward projection amplifies further. Policy "
         "interventions that improve demand visibility increase EWMA responsiveness."),
    ]

    for code, title, color, description in LOOPS:
        p = doc.add_paragraph()
        run = p.add_run(f"{code}  {title}")
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(*bytes.fromhex(color.lstrip("#")))
        _body(doc, description)
        doc.add_paragraph()

    # ── SD stock table ────────────────────────────────────────────────────────
    _heading(doc, "A.2  SD State Variables", 2)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    _table_header_row(tbl, ["Variable", "Unit", "2023 Initial Value", "Source"])
    SD_STOCKS = [
        ("Lithium stock",         "wks of consumption", "4.0",  "IEA CMMR 2023"),
        ("Cobalt stock",          "wks of consumption", "6.0",  "USGS MCS 2024"),
        ("Graphite stock",        "wks of consumption", "4.0",  "USGS MCS 2024"),
        ("REE / NdFeB stock",     "wks of consumption", "8.0",  "JOGMEC 2023"),
        ("SiC wafer stock",       "wks of consumption", "12.0", "Yole 2023"),
        ("Copper stock",          "wks of consumption", "3.0",  "ICSG 2023"),
        ("Cell inventory",        "GWh",                "63.2", "IEA GEO 2024 (4 wks × 15.81 GWh/wk)"),
        ("Cell capacity",         "GWh/yr",             "1 500","BNEF 2023 nameplate"),
        ("Chemistry mix (LFP)",   "fraction [0, 1]",    "0.403","IEA GEO 2024 + BNEF 2023"),
        ("EV demand",             "GWh/yr annualised",  "21.4", "IEA GEO 2024 UK focus"),
        ("OEM backlog",           "k vehicles",         "0",    "Initial condition"),
        ("Cobalt price index",    "dimensionless",      "1.0",  "2023 = 100 baseline"),
        ("Composite price signal","dimensionless",      "1.0",  "2023 = 100 baseline"),
    ]
    for row in SD_STOCKS:
        _add_row(tbl, row)
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # PART B — ODD PROTOCOL
    # ══════════════════════════════════════════════════════════════════════════
    _heading(doc, "Part B — ODD Protocol for the Agent-Based Model", 1)
    _body(doc,
        "This section describes the agent-based component of the hybrid model using the "
        "ODD (Overview, Design concepts, Details) protocol (Grimm et al., 2006, 2010, 2020). "
        "The ABM is tightly coupled with the SD layer described in Part A.")

    doc.add_paragraph()
    fig_arch = build_abm_arch()
    _fig_to_docx(doc, fig_arch, width_inches=6.2,
                 caption="Figure 2. Four-tier ABM agent architecture. "
                         "Each tier contains three behavioural archetypes. "
                         "The SD coupling arrow (right) indicates bidirectional "
                         "information exchange each time-step.")

    # ── O: Overview ───────────────────────────────────────────────────────────
    _heading(doc, "B.1  Overview", 2)

    _heading(doc, "B.1.1  Purpose", 3)
    _body(doc,
        "The model investigates how supply disruptions at different tiers of the global "
        "EV supply chain propagate to UK vehicle output, and how UK government policy "
        "interventions can dampen these effects. The primary outputs are weekly UK OEM "
        "production (k vehicles/week), cumulative production loss relative to baseline, "
        "order backlog, and commodity price indices over a 5-year (260-week) horizon.")

    _heading(doc, "B.1.2  Entities, State Variables, and Scales", 3)
    _body(doc, "The model contains five agent types across four supply chain tiers:")

    tbl2 = doc.add_table(rows=1, cols=5)
    tbl2.style = "Table Grid"
    _table_header_row(tbl2, ["Agent Type", "# Agents", "Key State Variables",
                              "Archetypes", "Tier"])
    AGENTS = [
        ("MineralSupplierAgent", "14",
         "output_fraction, weekly_supply_contribution, capacity, is_shocked",
         "StateBacked, WesternMiner, GreenfieldBuilder", "0"),
        ("CellManufacturerAgent", "9",
         "inventory_gwh, weekly_capacity, lfp_share, output_gwh",
         "PlatformLeader, HyperScaleChallenger, IncumbentUnderPressure", "1"),
        ("Tier1SupplierAgent", "4",
         "inventory, pipeline[], weekly_capacity, dual_source_active",
         "BatteryPackIntegrator, PremiumPowerElectronics, EstablishedVolumeSupplier", "2"),
        ("OEMAgent", "7",
         "production_k, target_k, halt_weeks, backlog_k, ice_fallback_fraction",
         "TransitioningLegacyOEM, EVNativeScaleAspirant, ProfitableEstablishedOEM", "3"),
        ("MarketAgent", "6",
         "demand_gwh, price_level, availability, backlog_sensitivity",
         "—", "Market"),
    ]
    for row in AGENTS:
        _add_row(tbl2, row)
    doc.add_paragraph()

    _body(doc,
        "Time step: 1 week (discrete Euler integration). Simulation horizon: 260 weeks "
        "(5 years). Spatial scale: global supply chain with UK-focused output tracking. "
        "Market scope fraction: ~0.024 (UK represents ≈2.4 % of global EV cell demand, "
        "used to scale mineral inflows from the global agent population to UK consumption).")

    _heading(doc, "B.1.3  Process Overview and Scheduling", 3)
    _body(doc, "Each weekly time-step executes in the following fixed sequence:")
    steps = [
        ("1", "SD pre-step",
         "SD.compute_input_fractions() — reads measured mineral stocks and component "
         "inventories; computes availability fractions and price indices exposed to agents."),
        ("2", "Shock application",
         "Active shocks modify agent shock_multiplier values for the current week."),
        ("3", "Policy ramp",
         "Policy parameters linearly ramp over 13 weeks from baseline to target values."),
        ("4", "Mineral agents step",
         "Each MineralSupplierAgent calls _compute_output_fraction() and reports "
         "weekly_supply_contribution (scaled by market_scope_fraction)."),
        ("5", "Cell agents step",
         "Each CellManufacturerAgent applies the Leontief production constraint over "
         "(Li, Co, graphite) inputs and updates inventory."),
        ("6", "Tier-1 agents step",
         "Each Tier1SupplierAgent updates the FIFO order pipeline, receives deliveries, "
         "and places new orders via _order_quantity()."),
        ("7", "OEM agents step",
         "Each OEMAgent applies the Leontief constraint over (pack, inverter, motor, "
         "harness) inputs and records production and backlog."),
        ("8", "Market agents step",
         "Each MarketAgent updates demand incorporating price elasticity, backlog "
         "sensitivity, and availability floor."),
        ("9", "SD post-step",
         "SD.update(flows) — absorbs aggregated ABM outputs; advances transit pipelines; "
         "updates perceived stocks; runs all five feedback loops (F1–F5)."),
        ("10", "Results recording",
         "Model records all state variables for output."),
    ]
    tbl3 = doc.add_table(rows=1, cols=3)
    tbl3.style = "Table Grid"
    _table_header_row(tbl3, ["Step", "Sub-model", "Description"])
    for num, sub, desc in steps:
        _add_row(tbl3, [num, sub, desc])
    doc.add_paragraph()

    # ── D: Design concepts ────────────────────────────────────────────────────
    _heading(doc, "B.2  Design Concepts", 2)

    concepts = [
        ("Basic principles",
         "The model is grounded in supply chain management theory (order-up-to policy, "
         "bullwhip effect, safety-stock targeting), industrial economics (Leontief "
         "production functions, price formation), and system dynamics (stock-flow "
         "structures, feedback loops). Archetype heterogeneity reflects firm-level "
         "strategic positioning documented in industry literature (IEA, BNEF, SNE Research)."),

        ("Emergence",
         "Supply disruption amplitudes and recovery timescales emerge from the interaction "
         "of agent ordering rules, inventory buffers, transit delays, and SD feedback loops. "
         "Price responses, chemistry substitution rates, and backlog accumulation are not "
         "prescribed but emerge from the stock-flow dynamics and agent decisions."),

        ("Adaptation",
         "Agents adapt to their environment through three mechanisms: (1) MineralSupplierAgent "
         "recovery rate adjusts gradually after a shock resolves; (2) CellManufacturerAgent "
         "LFP share shifts in response to cobalt price via the F2 loop; "
         "(3) TransitioningLegacyOEM activates ICE fallback when the margin signal "
         "drops below 0.85, and reverses at half the ramp rate (asymmetric adaptation)."),

        ("Objectives",
         "Agents do not optimise an explicit objective function. Instead, each archetype "
         "follows a satisficing heuristic encoded in its decision-rule hook. "
         "PlatformLeader targets a 70 % inventory cover before producing; "
         "ProfitableEstablishedOEM requires 50 % average inventory before running at full "
         "capacity; EVNativeScaleAspirant maximises growth within a 5 %/week boost ceiling."),

        ("Learning",
         "No explicit learning mechanism. Adaptation is reactive (stimulus–response) "
         "rather than anticipatory or evolutionary. The bullwhip smooth multiplier "
         "improves EWMA signal tracking under policy, which can be interpreted as "
         "improved demand visibility (digital supply chain) rather than learning."),

        ("Prediction",
         "PremiumPowerElectronics forward-projects demand 16 weeks ahead when calculating "
         "order quantities, representing the only anticipatory element in the model. "
         "All other agents are purely reactive to current stock levels and price signals."),

        ("Sensing",
         "Agents sense the SD model's measured inventory fractions and price indices "
         "(via model.get_input_fraction() and model.get_price_signal()). They do not "
         "observe other agents' states directly. Measurement lags mean agents see a "
         "smoothed (τ ≈ 2 week) version of true stock levels, not instantaneous values."),

        ("Interaction",
         "Agents interact indirectly through shared SD stocks: mineral agents write "
         "inflows; cell agents consume minerals and write cell inventory; Tier-1 agents "
         "consume cell output and write component inventories; OEM agents consume "
         "components and write vehicle output. There is no direct peer-to-peer "
         "communication between agents."),

        ("Stochasticity",
         "Weekly Gaussian noise (σ = 3–7 % by mineral) is applied to mineral supply "
         "flows. Seed 42 is used for all published results. Agent decision rules and "
         "shock parameters are deterministic given the seed."),

        ("Collectives",
         "Agents are grouped into archetypes (three per tier) and focus regions (UK, "
         "China, Germany/EU, US, Korea, Japan). The model can run in full-global mode "
         "(all 7 OEMs, 6 markets) or UK-focus mode (UK OEM and UK market only, with "
         "market_scope_fraction scaling applied to mineral inflows)."),

        ("Observation",
         "Weekly time-series are recorded for all state variables. The "
         "get_results() method returns a 260-row DataFrame. The validate_model.py "
         "suite runs 476 automated checks across 11 categories (static, invariants, "
         "baseline, scenario, propagation, policy, dynamics, consistency, real_timeseries, "
         "run, outputs) after each simulation."),
    ]

    for title_c, text in concepts:
        p = doc.add_paragraph()
        run = p.add_run(title_c + ".  ")
        run.font.bold = True
        run.font.size = Pt(10)
        p.add_run(text).font.size = Pt(10)
        p.paragraph_format.space_after = Pt(4)

    # ── D: Details ────────────────────────────────────────────────────────────
    _heading(doc, "B.3  Details", 2)

    _heading(doc, "B.3.1  Initialisation", 3)
    _body(doc,
        "The model is initialised from the 2023 calibration dataset. Mineral stocks "
        "are set to their target safety-stock levels (4–12 weeks depending on mineral). "
        "The SD mineral transit pipelines are pre-filled at steady-state "
        "(value = market_scope_fraction × transport_weeks) to eliminate the 8-week "
        "transient that would otherwise occur as the pipeline fills from zero. "
        "Cell inventory is set to 4 weeks of baseline throughput (63.2 GWh). "
        "Cell capacity is 1,500 GWh/yr (BNEF 2023). LFP share is 0.403 (IEA GEO 2024). "
        "OEM backlog is zero. All price indices are 1.0.")

    _heading(doc, "B.3.2  Input Data and Calibration", 3)
    tbl4 = doc.add_table(rows=1, cols=4)
    tbl4.style = "Table Grid"
    _table_header_row(tbl4, ["Parameter", "Value", "Source", "Calibration note"])
    CALIB = [
        ("UK OEM annual volume", "175 k vehicles/yr", "SMMT 2024",
         "JLR + MINI Oxford + Vauxhall Ellesmere Port"),
        ("UK market demand 2023", "20 GWh/yr", "IEA GEO 2024 UK",
         "Scaled from 712 GWh global EV energy"),
        ("UK demand growth rate", "28 %/yr", "Bloomberg NEF EVO 2024",
         "Compounded weekly: (1.28)^(1/52) − 1"),
        ("CATL global cell share", "37 %", "SNE Research 2023",
         "304 GWh of 822 GWh global production"),
        ("Cobalt transport delay", "8 weeks", "IEA CMMR 2023",
         "DRC mine → CMOC/Umicore → cell factory"),
        ("REE transport delay", "10 weeks", "JOGMEC 2023",
         "China mine → NdFeB magnet → motor"),
        ("Price adj. speed (upward)", "0.05 /wk", "Calibrated to Co spot",
         "5 % gap closed per week"),
        ("Price adj. speed (downward)", "0.075 /wk", "Calibrated to Co spot",
         "1.5× faster downward recovery"),
        ("LFP logistic midpoint", "1.30× cobalt baseline", "IEA CMMR 2023",
         "Switch point from industry survey"),
        ("Cell capex trigger", "85 % utilisation", "BNEF 2023",
         "Industry rule-of-thumb for gigafactory expansion"),
        ("Bullwhip factor", "1.25×", "Lee et al. 1997",
         "Empirical order amplification in electronics SC"),
        ("Price elasticity (UK)", "−0.36", "SMMT / BVRLA surveys",
         "EV demand sensitivity to price premium"),
        ("Market scope fraction", "≈0.024", "IEA GEO 2024",
         "UK active market GWh / global cell GWh"),
    ]
    for row in CALIB:
        _add_row(tbl4, row)
    doc.add_paragraph()

    _heading(doc, "B.3.3  Sub-model: Shock Mechanism", 3)
    _body(doc,
        "Shocks are parameterised as {target, start_week, end_week, severity}. "
        "The effective severity applied to an agent is:")
    p = doc.add_paragraph(
        "effective_severity = raw_severity × (1 − agent_absorption) × (1 − source_calibration) "
        "× (1 − policy_absorption)")
    p.runs[0].font.name = "Courier New"
    p.runs[0].font.size = Pt(9)
    _body(doc,
        "where agent_absorption is the archetype-specific resilience parameter "
        "(e.g., 0.15 for WesternMiner), source_calibration is derived from financial-profile "
        "data, and policy_absorption is added by active policy packages. "
        "During a shock, the agent's output_fraction = 1 − effective_severity. "
        "Recovery increments output_fraction by recovery_rate_wk each week until "
        "the shock end-week, after which full recovery is immediate.")

    _heading(doc, "B.3.4  Sub-model: Leontief Production (Cell & OEM)", 3)
    _body(doc,
        "Both CellManufacturerAgent and OEMAgent use a Leontief (perfect-complement) "
        "production function. For cells:")
    p = doc.add_paragraph(
        "cell_output = min(Li_fraction, Co_fraction, graphite_fraction, sic_fraction) "
        "× desired_production")
    p.runs[0].font.name = "Courier New"
    p.runs[0].font.size = Pt(9)
    _body(doc,
        "For OEMs, the binding constraint is the minimum of (pack, inverter, motor, harness) "
        "availability fractions times the weekly production target. This means a single "
        "scarce input halts all production — the structural fragility the model is designed "
        "to test. Chemistry mix modifies cobalt sensitivity: a higher LFP share reduces "
        "the weight of cobalt in the Leontief constraint.")

    _heading(doc, "B.3.5  Sub-model: Order-Up-To Inventory Policy (Tier 1)", 3)
    _body(doc,
        "Tier-1 agents follow an order-up-to policy with a FIFO delivery pipeline. "
        "Each week, the agent computes:")
    p = doc.add_paragraph(
        "order = max(0, target_inventory − on_hand − pipeline_total) × bullwhip_factor")
    p.runs[0].font.name = "Courier New"
    p.runs[0].font.size = Pt(9)
    _body(doc,
        "The FIFO pipeline has length = lead_time_weeks (2–16 weeks depending on "
        "archetype). Dual sourcing activates when inventory < 20 % of target, at a "
        "20 % cost premium (which raises the component price signal). "
        "EstablishedVolumeSupplier reduces orders to 80 % of standard when inventory > "
        "150 % of target (production smoothing).")

    _heading(doc, "B.3.6  Sub-model: Policy Packages", 3)
    _body(doc,
        "Three policy packages are evaluated, all linearly ramped in over 13 weeks:")

    POLICIES = [
        ("Tier-1 Resilience Package",
         "Raises Tier-1 safety stock (harness: 2→4 wk, inverter: 8→12 wk, motor: 6→10 wk, "
         "pack: 3→5 wk). Reduces harness lead time from 6 to 4 weeks. Raises UK OEM "
         "safety stock from 5 to 8 weeks. Adds 15 % shock absorption to Tier-1 agents."),
        ("Critical Minerals Security Package",
         "Adds 4-week strategic buffer to cobalt, graphite, REE, SiC stocks. Boosts "
         "mineral supply by 10–30 % via offtake/recycling proxy. Adds 15 % shock "
         "absorption to mineral agents. Raises lithium reserve to 6 weeks."),
        ("Full Industrial Strategy Package",
         "Combines both packages above. Additionally: raises UK OEM vertical integration "
         "from 5 % to 20 %; boosts UK demand growth rate by 5 pp; raises cell capacity "
         "growth rate; adds 20 % shock absorption across all tiers; reduces UK OEM "
         "target lead time by 2 weeks."),
    ]
    for pkg_name, pkg_desc in POLICIES:
        p = doc.add_paragraph()
        p.add_run(pkg_name + ".  ").font.bold = True
        p.add_run(pkg_desc)
        p.paragraph_format.space_after = Pt(4)

    _heading(doc, "B.3.7  Validation", 3)
    _body(doc,
        "The model is validated via an automated suite of 476 checks across 11 categories "
        "run by validate_model.py after each simulation (34 scenarios × 260 weeks, seed 42). "
        "All 476 checks pass (0 WARN, 0 FAIL). Categories include:")
    val_cats = [
        ("static",         "180 checks", "Parameter bounds, type checks, sensible defaults"),
        ("invariants",     "170 checks", "Non-negativity, conservation, monotonicity over time"),
        ("run",            " 34 checks", "Each scenario completes without error"),
        ("outputs",        " 34 checks", "Production > 0, no NaN values"),
        ("scenario",       " 14 checks", "Shock scenarios differ meaningfully from baseline"),
        ("propagation",    " 12 checks", "Mineral stock drop → price rise → cell → OEM cascade"),
        ("policy",         "  9 checks", "Policy scenarios improve mean OEM production"),
        ("consistency",    "  9 checks", "Cross-scenario ordering, policy > shock-only"),
        ("dynamics",       "  7 checks", "Demand growth, LFP bounds, bullwhip > 1"),
        ("baseline",       "  5 checks", "2023 anchor: production, stocks, price indices"),
        ("real_timeseries","  2 checks", "MAE against ONS/SMMT UK car production index"),
    ]
    tbl5 = doc.add_table(rows=1, cols=3)
    tbl5.style = "Table Grid"
    _table_header_row(tbl5, ["Category", "Count", "Description"])
    for cat, count, desc in val_cats:
        _add_row(tbl5, [cat, count, desc])

    doc.add_paragraph()
    _body(doc,
        "Real historical validation benchmarks the modelled UK OEM production against the "
        "ONS/SMMT seasonally adjusted UK car production index (January 2023 – October 2025, "
        "34 months). Both series are normalised to 2023 average = 100. "
        "The baseline MAE is 51.0 index points (MAPE 64.2 %), reflecting the expected "
        "divergence between a stylised EV-focused simulation and total UK car production "
        "(which includes ICE vehicles and a post-COVID recovery ramp not modelled here).")

    # ── References ─────────────────────────────────────────────────────────────
    _heading(doc, "References", 1)
    refs = [
        "BNEF (2023). Electric Vehicle Outlook 2023. BloombergNEF.",
        "Grimm, V., Berger, U., Bastiansen, F., et al. (2006). A standard protocol for "
        "describing individual-based and agent-based models. Ecological Modelling, 198, 115–126.",
        "Grimm, V., Berger, U., DeAngelis, D.L., et al. (2010). The ODD protocol: A review "
        "and first update. Ecological Modelling, 221, 2760–2768.",
        "Grimm, V., et al. (2020). The ODD protocol for describing agent-based and other "
        "simulation models: A second update to improve clarity, replication, and structural "
        "realism. Journal of Artificial Societies and Social Simulation, 23(2), 7.",
        "IEA (2024). Global EV Outlook 2024. International Energy Agency.",
        "IEA (2023). Critical Minerals Market Review 2023. International Energy Agency.",
        "Lee, H.L., Padmanabhan, V., Whang, S. (1997). The bullwhip effect in supply chains. "
        "Sloan Management Review, 38(3), 93–102.",
        "SMMT (2024). Electric Vehicle Market Statistics. Society of Motor Manufacturers "
        "and Traders.",
        "USGS (2024). Mineral Commodity Summaries 2024. U.S. Geological Survey.",
        "Yole Group (2023). Power SiC 2023: Materials, Devices and Applications Report.",
    ]
    for r in refs:
        p = _bullet(doc, r)
        p.runs[0].font.size = Pt(9)

    # ── Save ──────────────────────────────────────────────────────────────────
    out_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "results", "EV_Supply_Chain_CLD_ODD.docx"
    )
    doc.save(out_path)
    print(f"Saved: {out_path}")
    return out_path


if __name__ == "__main__":
    build_doc()
