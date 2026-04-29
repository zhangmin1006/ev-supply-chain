"""
Generate standardized reporting protocols for the EV supply-chain ABM+SD model.

Output:
  ABM_SD_Standardized_Reporting_Protocols.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from model.config import CELL_MAKERS, MARKETS, MINERALS, OEMS, TIER1
from model.sd_model import BASELINE_WK, TARGET_WEEKS
from model.shocks import SCENARIOS


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "ABM_SD_Standardized_Reporting_Protocols.docx"


def add_heading(document: Document, text: str, level: int) -> None:
    document.add_heading(text, level=level)


def add_para(document: Document, text: str = "", bold_prefix: str | None = None) -> None:
    paragraph = document.add_paragraph()
    if bold_prefix:
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        paragraph.add_run(text)
    else:
        paragraph.add_run(text)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_table(document: Document, headers: list[str], rows: list[list[object]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = "" if value is None else str(value)

    document.add_paragraph()


def setup_styles(document: Document) -> None:
    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Calibri"


def build_document() -> Document:
    document = Document()
    setup_styles(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Standardized Reporting Protocols\nEV Supply Chain ABM + SD Model")
    title_run.bold = True
    title_run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Agent-Based Model reporting protocol and System Dynamics stock-flow reporting protocol")

    add_heading(document, "1. Purpose and Scope", 1)
    add_para(
        document,
        "This document defines a standardized reporting protocol for the hybrid EV supply-chain model. "
        "It is intended to make simulation runs reproducible, comparable across scenarios, and auditable "
        "when model parameters are updated from firm-level, financial, or supply-chain data.",
    )
    add_bullets(
        document,
        [
            "ABM scope: heterogeneous agents for mineral sources, cell manufacturers, Tier-1 suppliers, OEM groups, and regional markets.",
            "SD scope: aggregate stock-and-flow layer for critical minerals, cells, and vehicle-equivalent subsystem inventories.",
            "Temporal unit: one simulation step equals one week.",
            "Default horizon: 260 weeks, equal to five years.",
            "Standard output location: results/*.csv plus comparison plots and summary workbooks.",
        ]
    )

    add_heading(document, "2. ABM Standardized Reporting Protocol", 1)
    add_heading(document, "2.1 Required Model Metadata", 2)
    add_table(
        document,
        ["Field", "Required Reporting Content"],
        [
            ["Model name", "EV Supply Chain ABM + SD Simulation"],
            ["Model version/date", "Report code version, run date, and any data workbook versions used."],
            ["Random seed", "Default 42 unless otherwise specified."],
            ["Time step", "Weekly."],
            ["Simulation horizon", "Number of weeks, normally 260."],
            ["Scenario set", "List all scenario IDs used from model.shocks.SCENARIOS."],
            ["Financial calibration", "State whether listed-company financial profiles were used and name the calibration workbook."],
        ],
    )

    add_heading(document, "2.2 Agent Classes and Reporting Fields", 2)
    add_table(
        document,
        ["Agent class", "Population", "Core state variables", "Required outputs"],
        [
            [
                "MineralSupplierAgent",
                "Country/source cohorts for lithium, cobalt, graphite, REE, and SiC wafer supply.",
                "mineral, country, global_share, shock_multiplier, recovery_rate_wk, financial_profile.",
                "weekly_supply_contribution, output_fraction, shock state.",
            ],
            [
                "CellManufacturerAgent",
                "Battery cell firms or firm cohorts.",
                "capacity, market_share, LFP/NMC mix, inventory_gwh, backlog_gwh, shock_multiplier.",
                "output_gwh, utilisation, inventory_gwh, backlog_gwh.",
            ],
            [
                "Tier1SupplierAgent",
                "Battery pack, inverter, motor, and harness subsystem cohorts.",
                "capacity_k, inventory, pipeline, key_input, input_dependency, dual_source_active.",
                "output_k, inventory, shortage, dual-source state.",
            ],
            [
                "OEMAgent",
                "Regional or firm-group OEM assemblers.",
                "weekly_target, component inventories, backlog_k, shock_multiplier.",
                "production_k, backlog_k, halt_weeks, average component inventory.",
            ],
            [
                "MarketAgent",
                "Regional end-market demand.",
                "gwh_annual, yoy_growth, avg_kwh, price_elasticity.",
                "weekly_demand_gwh, weekly_demand_k_veh.",
            ],
        ],
    )

    add_heading(document, "2.3 ABM Weekly Event Schedule", 2)
    add_numbered(
        document,
        [
            "Apply scheduled shocks and shock resolutions for the current week.",
            "Compute SD input availability fractions from current stock levels.",
            "Step mineral supplier agents and report mineral supply fractions.",
            "Step cell manufacturer agents and report cell production in GWh.",
            "Step Tier-1 supplier agents and report subsystem output in thousand vehicle-equivalents.",
            "Step OEM agents and report vehicle assembly output in thousand vehicles.",
            "Step market agents and update demand using growth and price response.",
            "Aggregate ABM outputs into SD inflow/outflow variables.",
            "Update SD stocks and record all model metrics.",
        ],
    )

    add_heading(document, "2.4 Financially Redesigned Agent Reporting", 2)
    add_para(
        document,
        "When listed-company financial data are available, each affected agent must report the financial peer companies "
        "used for calibration and the resulting behavioural multipliers.",
    )
    add_table(
        document,
        ["Multiplier", "Derived from", "Model effect", "Bounded interpretation"],
        [
            ["recovery_multiplier", "Operating margin, free cash flow margin, debt-to-assets.", "Scales post-shock recovery_rate_wk.", "Higher value means faster recovery."],
            ["inventory_multiplier", "Free cash flow margin and leverage.", "Scales safety_stock_weeks.", "Higher value means stronger buffer capacity."],
            ["growth_multiplier", "Revenue CAGR, capex intensity, leverage.", "Scales weekly capacity growth.", "Higher value means faster capacity expansion."],
            ["shock_absorption", "Profitability, cash generation, leverage.", "Reduces effective shock severity.", "Maximum effect is intentionally capped."],
        ],
    )

    add_heading(document, "2.5 ABM Scenario Reporting Standard", 2)
    add_para(document, "Every scenario run should report the following fields.")
    add_table(
        document,
        ["Field", "Definition"],
        [
            ["scenario", "Scenario ID, e.g. baseline, drc_cobalt, china_catl_disruption."],
            ["shock target", "Agent ID receiving the shock."],
            ["start_week / end_week", "Shock start and resolution week."],
            ["severity", "Fraction of output lost before financial shock absorption."],
            ["effective severity", "Severity after agent-level financial shock absorption."],
            ["peak loss", "Maximum production loss relative to baseline."],
            ["recovery week", "Last week below 90% of baseline plus one."],
            ["cumulative loss", "Sum of weekly production shortfall relative to baseline."],
        ],
    )

    add_heading(document, "3. SD Stock-and-Flow Reporting Protocol", 1)
    add_heading(document, "3.1 Stock Definitions", 2)
    add_table(
        document,
        ["Stock", "Unit", "Target weeks", "Baseline weekly throughput", "Interpretation"],
        [
            [stock, "normalised fraction" if stock in {"lithium", "cobalt", "graphite", "ree", "sic_wafer"} else ("GWh" if stock == "cells" else "k vehicle-equivalents"), TARGET_WEEKS[stock], BASELINE_WK[stock], "Critical input inventory held by the SD layer."]
            for stock in TARGET_WEEKS
        ],
    )

    add_heading(document, "3.2 Stock Update Equation", 2)
    add_para(document, "For every stock i and week t, the SD layer applies:")
    add_para(document, "Stock_i(t+1) = max(0, Stock_i(t) + Inflow_i(t) - Outflow_i(t))")
    add_para(document, "Stocks are capped at four times target inventory to prevent unbounded accumulation.")

    add_heading(document, "3.3 Required Flow Reporting", 2)
    add_table(
        document,
        ["Flow category", "Inflow definition", "Outflow definition"],
        [
            ["Mineral stocks", "Sum of mineral source output fractions by mineral.", "Cell, motor, or inverter production expressed as fraction of baseline mineral demand."],
            ["Cells", "Total cell-maker GWh output per week.", "Battery pack production converted to GWh demand."],
            ["Packs", "Battery pack Tier-1 output in k units.", "Total OEM vehicle production in k units."],
            ["Inverters", "Inverter Tier-1 output in k units.", "Total OEM vehicle production in k units."],
            ["Motors", "Motor Tier-1 output in k units.", "Total OEM vehicle production in k units."],
            ["Harness", "Harness Tier-1 output in k units.", "Total OEM vehicle production in k units."],
        ],
    )

    add_heading(document, "3.4 Input Availability Fractions", 2)
    add_para(document, "The SD layer exposes an input availability fraction to agents before each weekly step:")
    add_para(document, "input_fraction_i(t) = min(2.0, Stock_i(t) / TargetStock_i)")
    add_bullets(
        document,
        [
            "input_fraction < 1.0 indicates shortage pressure.",
            "input_fraction = 1.0 indicates target stock is available.",
            "input_fraction > 1.0 indicates surplus, capped at 2.0.",
            "Agents use these values in Leontief or partial-dependency production constraints.",
        ],
    )

    add_heading(document, "3.5 SD Output Metrics", 2)
    add_table(
        document,
        ["Metric", "Definition", "Required unit"],
        [
            ["stock_<name>_wk", "Weeks of baseline supply remaining for each stock.", "weeks"],
            ["cell_production_gwh", "Aggregate weekly cell production.", "GWh/week"],
            ["t1_<component>_k", "Weekly Tier-1 subsystem output.", "k units/week"],
            ["oem_production_k", "Aggregate weekly vehicle output.", "k vehicles/week"],
            ["market_demand_gwh", "Aggregate weekly battery demand.", "GWh/week"],
            ["price_signal", "Weighted mineral scarcity price index.", "index, baseline = 1"],
        ],
    )

    add_heading(document, "4. Parameter Reporting Requirements", 1)
    add_heading(document, "4.1 Current Model Population Counts", 2)
    add_table(
        document,
        ["Parameter group", "Count", "Source object"],
        [
            ["Minerals", len(MINERALS), "model.config.MINERALS"],
            ["Cell makers", len(CELL_MAKERS), "model.config.CELL_MAKERS"],
            ["Tier-1 subsystems", len(TIER1), "model.config.TIER1"],
            ["OEM groups", len(OEMS), "model.config.OEMS"],
            ["Markets", len(MARKETS), "model.config.MARKETS"],
            ["Scenarios", len(SCENARIOS), "model.shocks.SCENARIOS"],
        ],
    )

    add_heading(document, "4.2 Minimum Parameter Table for Publications", 2)
    add_bullets(
        document,
        [
            "Report all capacity values and units.",
            "Report market shares and source shares separately.",
            "Report chemistry mix for battery cell makers.",
            "Report input dependencies for Tier-1 subsystems, especially SiC and REE dependency fractions.",
            "Report safety stock weeks and recovery rates before and after financial calibration.",
            "Report all shock severity assumptions and source rationale.",
        ]
    )

    add_heading(document, "5. Reproducibility Checklist", 1)
    add_numbered(
        document,
        [
            "Record git commit hash or archive copy of the model code.",
            "Record Python version and package versions.",
            "Record random seed.",
            "Record scenario list and run horizon.",
            "Archive input workbooks, especially listed_company_tiers_financials.xlsx and financial_agent_calibration.xlsx.",
            "Export raw weekly outputs for every scenario.",
            "Export summary metrics and plots using the same baseline reference.",
            "Document any manual data changes or ticker substitutions.",
        ],
    )

    add_heading(document, "6. Recommended Reporting Tables", 1)
    add_table(
        document,
        ["Table", "Purpose"],
        [
            ["Agent inventory", "Lists all agents, class, tier, region, capacity, financial peers, and multipliers."],
            ["Stock-flow inventory", "Lists every SD stock, unit, target, baseline throughput, inflow, and outflow."],
            ["Scenario design", "Lists shock target, timing, severity, and rationale."],
            ["Weekly outputs", "Reports model time series for production, demand, stocks, price signal, and backlog."],
            ["Scenario summary", "Reports peak loss, mean loss, weeks below threshold, recovery week, and cumulative loss."],
            ["Calibration audit", "Reports data source, source year, transformation rule, and final parameter value."],
        ],
    )

    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    return document


def main() -> None:
    document = build_document()
    document.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
