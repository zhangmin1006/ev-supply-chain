"""
Generate Word documentation for the EV Supply Chain Hybrid ABM+SD Model.
Covers:
  - ODD Protocol (Overview, Design Concepts, Details)
  - System Dynamics model description
  - Causal Loop Diagram narrative
  - Agent decision rules and calibration
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__),
                           "EV_Supply_Chain_Model_Documentation.docx")

# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_para(doc, text, bold=False, italic=False, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(indent * 0.4)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.add_run(text)
    return p

def add_numbered(doc, text, level=0):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.add_run(text)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), "D9E2F3")
        cell._tc.get_or_add_tcPr().append(shading)
    # Data rows
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            cell.text = str(val)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Column widths
    if col_widths:
        for col_idx, w in enumerate(col_widths):
            for cell in table.columns[col_idx].cells:
                cell.width = Inches(w)
    return table

def add_equation_box(doc, text):
    """Add a grey-shaded equation / formula block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)
    # Shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p

def set_doc_margins(doc, top=1.0, bottom=1.0, left=1.25, right=1.25):
    section = doc.sections[0]
    section.top_margin    = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin   = Inches(left)
    section.right_margin  = Inches(right)

# ─────────────────────────────────────────────────────────────────────────────
# Build document
# ─────────────────────────────────────────────────────────────────────────────

doc = Document()
set_doc_margins(doc)

# ── Title ────────────────────────────────────────────────────────────────────
title = doc.add_heading("EV Supply Chain Hybrid ABM + SD Model", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Technical Documentation — ODD Protocol, System Dynamics Design, "
                "Causal Loop Diagram, and Agent Specification")
r.italic = True
r.font.size = Pt(11)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run("Queen's University Belfast  ·  Model Building Project  ·  2024–2025").font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Table of Contents", 1)
toc_items = [
    "1.  Introduction and Model Purpose",
    "2.  ODD Protocol",
    "    2.1  Overview",
    "    2.2  Design Concepts",
    "    2.3  Details",
    "3.  System Dynamics (SD) Layer",
    "    3.1  Architecture and Stocks",
    "    3.2  Feedback Loops",
    "    3.3  Causal Loop Diagram",
    "    3.4  Stock–Flow Equations",
    "    3.5  Parameterisation",
    "4.  Agent-Based Model (ABM) Layer",
    "    4.1  Mineral Supplier Agent",
    "    4.2  Cell Manufacturer Agent",
    "    4.3  Tier-1 Supplier Agent",
    "    4.4  OEM Agent",
    "    4.5  Market Agent",
    "5.  Hybrid Coupling: ABM ↔ SD",
    "6.  Financial Calibration from Listed Companies",
    "7.  Scenario Shocks",
    "8.  Parameter Summary",
    "9.  References",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Inches(0.25 if item.startswith("    ") else 0)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1.  Introduction and Model Purpose", 1)

add_para(doc,
    "The global electric vehicle (EV) supply chain is characterised by high geographic "
    "concentration, complex multi-tier interdependencies, and deep vulnerability to "
    "commodity price volatility. This model provides a computational framework for "
    "simulating disruption propagation, price dynamics, and resilience across four "
    "supply-chain tiers — from raw material extraction through to final vehicle assembly.")

add_para(doc,
    "The model integrates two complementary paradigms:")
add_bullet(doc,
    "Agent-Based Modelling (ABM): captures heterogeneous firm behaviour, inventory "
    "management decisions, shock absorption capacity, and emergent supply-chain dynamics "
    "arising from decentralised agent interactions.")
add_bullet(doc,
    "System Dynamics (SD): provides aggregate stock-flow-feedback structure for "
    "commodity prices, chemistry mix, capacity investment cycles, and demand dynamics "
    "that operate at the industry level above individual firm decisions.")

add_para(doc,
    "The four-tier supply chain framework is calibrated using financial and operational "
    "data from 61 listed companies across all tiers, retrieved from Yahoo Finance. "
    "The simulation horizon is five years (260 weeks) at weekly time steps.")

doc.add_paragraph()
add_para(doc, "Supply Chain Tiers Modelled:", bold=True)
add_table(doc,
    ["Tier", "Label", "Agents / Firms", "Key Products"],
    [
        ["Tier 1", "Raw & Processed Materials",
         "14 mineral supplier agents across 5 minerals",
         "Lithium, cobalt, graphite, rare earth elements, SiC wafers"],
        ["Tier 2", "Core Components",
         "9 cell manufacturer agents",
         "LFP and NMC/NCA battery cells (GWh)"],
        ["Tier 3", "Subsystems / Tier-1 Integration",
         "4 subsystem supplier agents",
         "Battery packs, inverters, motors, wiring harnesses"],
        ["Tier 4", "OEM Assembly & Vehicle Integration",
         "7 OEM agents (by region)",
         "Assembled electric vehicles (k units/week)"],
    ],
    col_widths=[0.6, 1.8, 2.2, 2.4]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. ODD PROTOCOL
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2.  ODD Protocol", 1)
add_para(doc,
    "The ODD (Overview, Design concepts, Details) protocol (Grimm et al., 2006; 2010; 2020) "
    "is the standard framework for describing agent-based models. Each section below follows "
    "the ODD format to ensure the model is fully reproducible and transparent.")

# ── 2.1 Overview ──────────────────────────────────────────────────────────────
add_heading(doc, "2.1  Overview", 2)

add_heading(doc, "2.1.1  Purpose", 3)
add_para(doc,
    "The model is designed to answer three research questions:")
add_numbered(doc,
    "How do supply shocks at Tier 1 (e.g., DRC cobalt disruption, China REE export "
    "restriction) propagate through the supply chain to affect OEM vehicle production?")
add_numbered(doc,
    "How do endogenous feedback mechanisms — commodity price formation, chemistry "
    "substitution, and capacity investment — moderate or amplify disruption impacts "
    "over a 5-year horizon?")
add_numbered(doc,
    "How do firm-level financial characteristics (balance sheet resilience, inventory "
    "policies, recovery capacity) determine heterogeneous responses to the same shock?")

add_heading(doc, "2.1.2  Entities, State Variables, and Scales", 3)

add_para(doc, "Agents:", bold=True)
add_table(doc,
    ["Agent Type", "Count", "Key State Variables", "Tier"],
    [
        ["MineralSupplierAgent", "14",
         "output_fraction, shock_multiplier, mineral, country, global_share",
         "Tier 1"],
        ["CellManufacturerAgent", "9",
         "weekly_capacity (GWh/wk), inventory_gwh, output_gwh, lfp_fraction, nmc_fraction",
         "Tier 2"],
        ["Tier1SupplierAgent", "4",
         "inventory (k veh-equiv), pipeline (order queue), output_k, dual_source_active",
         "Tier 3"],
        ["OEMAgent", "7",
         "inv{packs, inverters, motors, harness}, production_k, backlog_k, halt_weeks",
         "Tier 4"],
        ["MarketAgent", "6",
         "weekly_demand_gwh, _trend_demand_gwh, yoy_growth, price_elasticity",
         "Demand"],
    ],
    col_widths=[1.6, 0.6, 3.2, 0.9]
)

doc.add_paragraph()
add_para(doc, "SD Model State Variables (stocks):", bold=True)
add_table(doc,
    ["Variable", "Unit", "Initial Value", "Description"],
    [
        ["stocks[lithium]",   "wks consumption", "4.0",  "Lithium inventory in weeks of EV-industry consumption"],
        ["stocks[cobalt]",    "wks consumption", "6.0",  "Cobalt inventory"],
        ["stocks[graphite]",  "wks consumption", "4.0",  "Graphite anode inventory"],
        ["stocks[ree]",       "wks consumption", "8.0",  "Rare earth element inventory"],
        ["stocks[sic_wafer]", "wks consumption", "12.0", "Silicon carbide wafer inventory"],
        ["stocks[cells]",     "GWh",             "63.2", "Aggregate cell inventory (4 wks × 15.81 GWh/wk)"],
        ["stocks[packs]",     "k veh-equiv",     "807.7","Battery pack inventory"],
        ["stocks[inverters]", "k veh-equiv",     "2154", "Inverter inventory"],
        ["stocks[motors]",    "k veh-equiv",     "1615", "Motor inventory"],
        ["stocks[harness]",   "k veh-equiv",     "538.5","Wiring harness inventory (JIT)"],
        ["prices[mineral]",   "index (1.0=2023)","1.0",  "Commodity price index per mineral"],
        ["cell_capacity",     "GWh/yr",          "1500", "Global cell nameplate capacity"],
        ["cell_capacity_wip", "GWh/yr",          "0.0",  "Capacity under construction"],
        ["lfp_share",         "fraction [0,1]",  "0.403","LFP chemistry share of global cell output"],
        ["ev_demand_gwh_yr",  "GWh/yr",          "820",  "Annualised EV battery demand"],
        ["oem_backlog_k",     "k vehicles",      "0.0",  "Cumulative unfulfilled OEM orders"],
        ["bullwhip_index",    "dimensionless",   "1.0",  "EWMA order amplification ratio"],
    ],
    col_widths=[1.6, 1.4, 1.1, 2.9]
)

doc.add_paragraph()
add_para(doc, "Scales:", bold=True)
add_bullet(doc, "Temporal: discrete weekly time steps (dt = 1 week); horizon = 260 weeks (5 years).")
add_bullet(doc, "Spatial: global (no spatial topology); agents are distinguished by tier and region, not geography.")

add_heading(doc, "2.1.3  Process Overview and Scheduling", 3)
add_para(doc,
    "Each week, the following processes execute in strict sequence:")
add_numbered(doc, "Shock injection: any shocks scheduled for this week are applied to the targeted agents.")
add_numbered(doc, "SD → ABM signal: the SD model computes input_fractions (availability [0,2]) and "
             "price_signals for all stocks; these are read by agents in steps 3–6.")
add_numbered(doc, "Mineral suppliers step: each MineralSupplierAgent updates its output_fraction "
             "based on shock state and recovery rate.")
add_numbered(doc, "Cell manufacturers step: each CellManufacturerAgent applies the Leontief "
             "production constraint, updates inventory, and fulfils downstream demand.")
add_numbered(doc, "Tier-1 suppliers step: each Tier1SupplierAgent receives pipeline deliveries, "
             "places new orders (order-up-to + bullwhip), and delivers from inventory.")
add_numbered(doc, "OEM agents step: each OEMAgent receives component deliveries, applies the "
             "Leontief production constraint across all four subsystems, and records shortfalls.")
add_numbered(doc, "Market agents step: each MarketAgent advances trend demand and applies "
             "price-level adjustment.")
add_numbered(doc, "ABM → SD aggregation: all agent outputs are aggregated into inflow/outflow "
             "pairs and passed to the SD model via _collect_flows().")
add_numbered(doc, "SD update: the SD model updates all stocks (Euler integration), runs all "
             "five feedback loops, and records state.")
add_numbered(doc, "Metrics recording: a full snapshot of all state variables is appended to "
             "the results DataFrame.")

# ── 2.2 Design Concepts ───────────────────────────────────────────────────────
add_heading(doc, "2.2  Design Concepts", 2)

add_heading(doc, "2.2.1  Basic Principles", 3)
add_para(doc,
    "The model is grounded in three theoretical frameworks:")
add_bullet(doc,
    "Leontief production functions: production at each tier is constrained by the "
    "scarcest input (min operator). This is consistent with the 'weakest link' "
    "vulnerability of just-in-time manufacturing supply chains.")
add_bullet(doc,
    "Order-up-to inventory policy: each agent maintains a target stock level and "
    "orders enough each period to restore the inventory position (on-hand + in-transit) "
    "to the target. This policy is known to generate bullwhip amplification (Lee et al., 1997).")
add_bullet(doc,
    "System dynamics feedback structure: aggregate industry-level dynamics (price "
    "formation, chemistry substitution, capacity investment) are modelled as stock-flow "
    "systems with explicit feedback loops rather than as agent-level decisions, reducing "
    "model complexity while preserving system-level behaviour.")

add_heading(doc, "2.2.2  Emergence", 3)
add_para(doc,
    "The following macro-level patterns emerge from micro-level agent interactions:")
add_bullet(doc,
    "Bullwhip effect: upstream order amplification arises naturally from decentralised "
    "order-up-to decisions without any global coordinator.")
add_bullet(doc,
    "Supply-demand price spirals: commodity price spikes emerge when multiple "
    "cell manufacturers simultaneously chase lithium or cobalt inventory.")
add_bullet(doc,
    "Chemistry substitution waves: industry-wide LFP adoption accelerates when "
    "cobalt prices exceed threshold levels, as both the SD chemistry stock and "
    "individual agent NMC fractions respond to the same price signal.")
add_bullet(doc,
    "Production halt cascades: a mineral shock at Tier 1 causes cell shortages "
    "at Tier 2, pack shortages at Tier 3, and OEM production halts at Tier 4 — "
    "the cascade depth and timing depend on buffer stock levels along the chain.")

add_heading(doc, "2.2.3  Adaptation", 3)
add_para(doc,
    "Agents adapt their behaviour in response to supply conditions:")
add_bullet(doc,
    "Dual sourcing: Tier1SupplierAgents and OEMAgents activate dual sourcing "
    "(paying a 20% cost premium) when inventory falls below 20% of target, "
    "increasing order rates to restore stocks faster.")
add_bullet(doc,
    "Chemistry adaptation: CellManufacturerAgents reduce their effective NMC "
    "fraction when cobalt prices are high, blending in more LFP sub-cells "
    "even before the industry-level LFP share stock adjusts.")
add_bullet(doc,
    "Capacity growth: all agents grow their production capacity at the IEA-calibrated "
    "29%/yr baseline rate, modulated by their financial profile's growth_multiplier.")

add_heading(doc, "2.2.4  Objectives", 3)
add_para(doc,
    "Agents pursue implicit objectives through their decision rules:")
add_bullet(doc, "Cell manufacturers: maximise output subject to input availability; "
           "maintain safety-stock target.")
add_bullet(doc, "Tier-1 suppliers: maintain inventory position at target; minimise "
           "shortage weeks.")
add_bullet(doc, "OEMs: maximise vehicle production up to 110% of weekly target; "
           "clear backlog at 15%/week.")
add_bullet(doc, "Market agents: no optimisation — demand is exogenous with "
           "price-level adjustment.")

add_heading(doc, "2.2.4a  Production, Inventory, and Capacity Decisions", 3)
add_para(doc,
    "Agent types differ in how they decide production, inventory, and capacity. "
    "The model intentionally avoids one universal firm rule because each supply-chain "
    "tier faces a different operational constraint: geological supply shares at Tier 1, "
    "chemistry-specific materials at Tier 2, lead-time ordering at Tier 3, assembly "
    "synchronisation at Tier 4, and price/backlog-sensitive demand in markets.")
add_table(doc,
    ["Tier / Agent Type", "Production Decision", "Inventory Decision", "Capacity Decision", "Interaction Rule"],
    [
        [
            "Tier 1: Mineral supplier agents",
            "Weekly supply contribution = global source share x shock multiplier; after a shock resolves, output recovers at the calibrated weekly recovery rate.",
            "No private firm inventory decision; mineral inventory is represented by SD stocks measured in weeks of EV-industry consumption.",
            "No agent-level investment decision; mineral supply expansion is handled by the SD mineral supply scale with mineral-specific growth rates.",
            "Supply enters mineral transit pipelines, then SD stocks; downstream agents read measured input_fractions.",
        ],
        [
            "Tier 2: LFP-heavy cell manufacturers",
            "Produce up to weekly capacity subject to lithium, graphite, shock state, and weak cobalt exposure because LFP cells require no cobalt.",
            "Order-up-to target inventory: platform leaders replenish faster than hyper-scale challengers and NMC/NCA incumbents.",
            "Weekly cell capacity grows at the baseline growth rate multiplied by the financial growth multiplier.",
            "Demand is allocated by cell-maker market share; output feeds cell inventory, pack production, and mineral consumption.",
        ],
        [
            "Tier 2: NMC/NCA-heavy cell manufacturers",
            "Same capacity and inventory rule as LFP-heavy makers, but production has stronger cobalt sensitivity through the NMC fraction.",
            "Higher safety stock for some makers; unfulfilled downstream demand becomes cell backlog.",
            "Capacity growth is financially modulated; the SD layer separately tracks aggregate cell-capacity investment delays.",
            "High cobalt price reduces effective NMC dependence and raises industry LFP share, creating chemistry adaptation.",
        ],
        [
            "Tier 3: Battery-pack supplier",
            "Pack output is constrained by capacity, pack inventory, and cell availability.",
            "Uses lead-time pipeline and order-up-to policy; shortfalls are amplified by the bullwhip factor.",
            "Weekly capacity grows with the demand trend and financial multiplier.",
            "Pack deliveries are allocated to OEMs by OEM production share.",
        ],
        [
            "Tier 3: Inverter supplier",
            "Inverter output is partially constrained by SiC wafer availability; non-SiC inverter production can continue during SiC shortage.",
            "Uses order-up-to inventory, long lead time, shortage tracking, and dual sourcing below the inventory trigger.",
            "Capacity grows weekly; recovery is slower because power electronics manufacturing is specialised.",
            "SiC shortages lower inverter deliveries and can constrain OEM assembly.",
        ],
        [
            "Tier 3: Motor supplier",
            "Motor output is partially constrained by REE availability for PMSM motors; high REE price reduces effective REE dependence through motor-design substitution.",
            "Uses order-up-to inventory and lead-time pipeline, recording unmet component demand as shortage.",
            "Capacity grows weekly with financial adjustment.",
            "REE shortage affects motor deliveries, which become a Leontief input to OEM assembly.",
        ],
        [
            "Tier 3: Wiring-harness supplier",
            "Harness output is mainly capacity and shock constrained; copper is tracked in SD but current ABM dependency is zero.",
            "Very low safety stock represents JIT harness supply; disruptions propagate quickly.",
            "Capacity grows weekly and recovery is comparatively faster after shock resolution.",
            "Harness output is required one-for-one by OEMs; shortage quickly reduces vehicle assembly.",
        ],
        [
            "Tier 4: OEM agents",
            "Vehicle output follows a Leontief component constraint, with vertical integration cushioning part of component shortages before the shock multiplier is applied.",
            "OEMs receive components, consume one of each per vehicle, accumulate backlog from shortfalls, and clear backlog using surplus production.",
            "Weekly assembly target grows over time and is financially modulated; recovery rate controls post-shock throughput restoration.",
            "OEMs pull component demand upstream and convert component shortages into production loss and backlog.",
        ],
        [
            "Markets: regional demand agents",
            "No production; realised demand is trend demand adjusted by battery price signal and backlog/availability pressure.",
            "No inventory; high backlog reduces realised near-term demand through region-specific deferral, cancellation, or switching sensitivity.",
            "No capacity decision; regional demand trend follows configured YoY growth.",
            "Demand drives cell demand and OEM production targets; price and backlog close feedback loops.",
        ],
    ],
    col_widths=[1.55, 1.7, 1.7, 1.55, 1.7]
)

add_heading(doc, "2.2.5  Learning", 3)
add_para(doc,
    "Agents do not learn in the machine-learning sense. Adaptive responses are "
    "hard-coded behavioural rules (dual sourcing triggers, chemistry adjustment) "
    "rather than learned strategies. This is consistent with empirical observations "
    "that supply chain resilience measures are policy-defined, not emergent from learning.")

add_heading(doc, "2.2.6  Prediction", 3)
add_para(doc,
    "Agents do not form explicit forecasts of future prices or supply conditions. "
    "The order-up-to policy uses current inventory position as the signal, "
    "effectively assuming demand and lead times remain constant — the standard "
    "assumption in the bullwhip literature that generates the most conservative "
    "(amplifying) ordering behaviour.")

add_heading(doc, "2.2.7  Sensing", 3)
add_para(doc,
    "Each agent has access to:")
add_bullet(doc,
    "Global signals: SD input_fractions and price_signals (shared state, "
    "representing publicly available market information such as commodity spot prices).")
add_bullet(doc,
    "Local state only: agents do not observe other agents' inventories or order "
    "quantities directly, consistent with real supply chains where such information "
    "is typically proprietary.")

add_heading(doc, "2.2.8  Interaction", 3)
add_para(doc,
    "Agent interactions are mediated through the model's query methods:")
add_bullet(doc,
    "get_cell_demand(name): returns the fraction of total weekly cell demand "
    "allocated to a specific cell maker based on market share.")
add_bullet(doc,
    "get_component_deliveries(oem_name): routes Tier-1 output to each OEM in "
    "proportion to annual production targets.")
add_bullet(doc,
    "Agents do not trade directly with each other; all flows are aggregated "
    "through the hybrid model's collect_flows() method.")

add_heading(doc, "2.2.9  Stochasticity", 3)
add_para(doc,
    "The baseline model is deterministic given a fixed random seed. Stochasticity "
    "may be introduced through:")
add_bullet(doc, "Shock timing and severity in scenario runs.")
add_bullet(doc, "Future extensions: stochastic demand shocks, random lead-time "
           "variation, and Monte Carlo resilience testing.")

add_heading(doc, "2.2.10  Collectives", 3)
add_para(doc,
    "Agents are not grouped into formal collectives. The OEM agent groups "
    "(Chinese OEMs, US OEMs, etc.) are separate agents sharing the same "
    "tier-level configuration rather than a collective entity.")

add_heading(doc, "2.2.11  Observation", 3)
add_para(doc,
    "At every time step, a full snapshot of 40+ state variables is recorded, including:")
add_bullet(doc, "ABM outputs: cell production GWh, OEM production k-vehicles per region, "
           "Tier-1 subsystem output, market demand, total backlog.")
add_bullet(doc, "SD stocks: weeks of supply for all 10 inventory stocks.")
add_bullet(doc, "SD prices: price indices for all 5 minerals.")
add_bullet(doc, "SD derived: LFP share, cell capacity utilisation, EV demand GWh/yr, "
           "OEM backlog, bullwhip index.")

# ── 2.3 Details ───────────────────────────────────────────────────────────────
add_heading(doc, "2.3  Details", 2)

add_heading(doc, "2.3.1  Initialisation", 3)
add_para(doc,
    "The model is initialised to the 2023 global EV supply chain baseline:")
add_bullet(doc, "Cell deployment: 822 GWh (IEA GEO 2024).")
add_bullet(doc, "Cell nameplate capacity: 1,500 GWh/yr (BNEF Battery Market Outlook 2023).")
add_bullet(doc, "EV production: 14,000 k vehicles (IEA GEO 2024).")
add_bullet(doc, "LFP share: 40.3% of deployed GWh (computed from cell maker chemistry profiles).")
add_bullet(doc, "All inventory stocks initialised at safety-stock target levels.")
add_bullet(doc, "All price indices initialised at 1.0 (2023 baseline).")
add_bullet(doc,
    "Financial calibration multipliers loaded from listed-company financial data "
    "(Yahoo Finance 5-year fundamentals); agents with no data use defaults of 1.0.")

add_heading(doc, "2.3.2  Input Data", 3)
add_para(doc, "Primary empirical sources:")
add_table(doc,
    ["Parameter Group", "Source", "Key Values"],
    [
        ["Mineral supply shares", "USGS MCS 2024",
         "DRC 70% cobalt; China 85% REE; Australia 46% lithium"],
        ["EV deployment & growth", "IEA GEO 2024",
         "822 GWh deployed 2023; 29%/yr growth"],
        ["Cell chemistry mix", "IEA GEO 2024; BNEF 2023",
         "LFP 40.3%; NMC/NCA 59.7%"],
        ["Battery pack price", "BNEF Battery Price Survey 2023",
         "~$139/kWh pack; price elasticity −0.30"],
        ["SiC wafer supply", "Yole Intelligence 2023",
         "Wolfspeed 30%; SiC 45% of EV inverters"],
        ["Price dynamics", "Benchmark Mineral Intelligence 2021–22",
         "Li₂CO₃ doubled in ~20 weeks → speed = 0.05/wk"],
        ["Gigafactory build time", "BNEF Gigafactory Tracker 2023",
         "2 years (104 weeks) from groundbreaking to production"],
        ["OEM production targets", "IEA GEO 2024; company AR 2023",
         "BYD 1,575k; CATL 37% cell share"],
        ["Financial calibration", "Yahoo Finance 5-yr fundamentals",
         "61 listed companies; 34 agent peer groups"],
        ["UK market", "SMMT 2024; UK DfT ZEV mandate",
         "315k BEV registrations; 28%/yr growth"],
    ],
    col_widths=[1.8, 2.0, 3.2]
)

add_heading(doc, "2.3.3  Submodels", 3)
add_para(doc,
    "Detailed submodel equations are given in Sections 3 (SD) and 4 (ABM). "
    "The key submodels are:")
add_bullet(doc,
    "Leontief production constraint (CellManufacturerAgent, OEMAgent): "
    "production = capacity × min(input_fractions).")
add_bullet(doc,
    "Order-up-to policy (Tier1SupplierAgent): "
    "order = weekly_target + bullwhip_factor × max(0, target_inventory − inventory_position).")
add_bullet(doc,
    "First-order price adjustment (SDModel._step_prices): "
    "p(t+1) = p(t) + speed × (1/availability − p(t)).")
add_bullet(doc,
    "Chemistry substitution (SDModel._step_chemistry_mix): "
    "piecewise-linear LFP target mapped from cobalt price; first-order convergence.")
add_bullet(doc,
    "Capacity investment pipeline (SDModel._step_cell_capacity): "
    "investment triggers at 85% utilisation; 104-week first-order delay.")
add_bullet(doc,
    "Battery price index (SDModel._battery_price_signal): "
    "bill-of-materials weighted average of mineral price indices.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 3. SYSTEM DYNAMICS LAYER
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3.  System Dynamics (SD) Layer", 1)

add_para(doc,
    "The SD layer provides aggregate industry-level stocks and feedback mechanisms "
    "that operate above individual firm decisions. It uses discrete-time Euler "
    "integration with dt = 1 week. The SD model is structured in four tiers "
    "mirroring the ABM agent tiers.")

# ── 3.1 Architecture ──────────────────────────────────────────────────────────
add_heading(doc, "3.1  Architecture and Stocks", 2)

add_para(doc,
    "Stocks are grouped by supply chain tier. All mineral stocks are measured "
    "in weeks of EV-industry baseline consumption, making shortage severity "
    "immediately interpretable regardless of the mineral's unit of measurement.")

add_para(doc, "Tier 1 — Raw & Processed Materials:", bold=True)
add_table(doc,
    ["Stock", "Unit", "Target Level", "Rationale for Target"],
    [
        ["Lithium",   "wks consumption", "4 weeks",  "IEA CMMR 2023: typical industry buffer"],
        ["Cobalt",    "wks consumption", "6 weeks",  "DRC political risk; USGS MCS 2024"],
        ["Graphite",  "wks consumption", "4 weeks",  "Standard anode supply buffer"],
        ["REE",       "wks consumption", "8 weeks",  "Slow processing; China-dominant supply"],
        ["SiC wafer", "wks consumption", "12 weeks", "26-week lead time; strategic buffer needed"],
    ],
    col_widths=[1.1, 1.4, 1.1, 3.4]
)
doc.add_paragraph()

add_para(doc, "Tier 2 — Core Components:", bold=True)
add_table(doc,
    ["Stock", "Unit", "Target Level", "Rationale"],
    [
        ["Cells",        "GWh",        "4 weeks × 15.81 GWh/wk = 63.2 GWh", "Standard cell inventory"],
        ["Cell capacity","GWh/yr",     "1,500 GWh/yr (2023 baseline)",       "BNEF nameplate estimate"],
        ["LFP share",    "fraction",   "0.403 (40.3%)",                      "IEA GEO 2024 chemistry mix"],
    ],
    col_widths=[1.2, 1.1, 2.6, 2.1]
)
doc.add_paragraph()

add_para(doc, "Tier 3 — Subsystems:", bold=True)
add_table(doc,
    ["Stock", "Target Level", "Note"],
    [
        ["Battery packs", "3 weeks", "Pack assembly"],
        ["Inverters",     "8 weeks", "SiC tension forces large buffer"],
        ["Motors",        "6 weeks", "REE magnet supply risk"],
        ["Wiring harness","2 weeks", "JIT — the critical vulnerability"],
    ],
    col_widths=[1.4, 1.4, 4.2]
)
doc.add_paragraph()

add_para(doc, "Tier 4 — OEM and Market:", bold=True)
add_table(doc,
    ["Stock", "Unit", "Initial Value", "Description"],
    [
        ["EV demand",    "GWh/yr",       "820",   "Annualised EV battery demand (IEA GEO 2024)"],
        ["OEM backlog",  "k vehicles",   "0",     "Cumulative unfulfilled vehicle orders"],
        ["Bullwhip index","dimensionless","1.0",  "Smoothed Tier-3/OEM order amplification ratio"],
    ],
    col_widths=[1.3, 1.3, 1.2, 3.2]
)

# ── 3.2 Feedback Loops ────────────────────────────────────────────────────────
add_heading(doc, "3.2  Feedback Loops", 2)

add_para(doc,
    "The SD model implements five feedback loops, each representing a distinct "
    "real-world market mechanism. Loops are labelled F1–F5.")

add_heading(doc, "F1 — Supply-Demand Price Formation Loop (Balancing)", 3)
add_para(doc,
    "When a mineral stock falls below its target level, the implied shortage "
    "drives commodity prices upward, which in turn reduces downstream demand "
    "and encourages supply expansion — a classic balancing (negative) feedback loop.")

add_para(doc, "Causal chain:", bold=True)
add_para(doc,
    "  Mineral stock ↓  →  availability fraction ↓  →  price ↑  →  "
    "cell cost ↑  →  battery price ↑  →  EV demand growth ↓  →  "
    "mineral outflow ↓  →  mineral stock stabilises")

add_equation_box(doc,
    "p(t+1) = p(t) + α × [ (1 / availability(t)) − p(t) ]\n"
    "\n"
    "where:  α = PRICE_ADJ_SPEED = 0.05 per week\n"
    "        availability = stock / (target_weeks × baseline_throughput)\n"
    "        price floor = 0.10,  price ceiling = 6.00  (× 2023 baseline)")

add_heading(doc, "F2 — Chemistry Substitution Loop (Balancing)", 3)
add_para(doc,
    "Rising cobalt prices create an economic incentive to shift battery chemistry "
    "from NMC (cobalt-intensive) toward LFP (cobalt-free). As LFP share increases, "
    "cobalt consumption falls, moderating the cobalt price — a balancing feedback loop.")

add_para(doc, "Causal chain:", bold=True)
add_para(doc,
    "  Cobalt price ↑  →  LFP target share ↑  →  LFP adoption ↑  →  "
    "cobalt consumption ↓  →  cobalt stock recovers  →  cobalt price ↓")

add_equation_box(doc,
    "LFP_target = 0.30 + t × (0.88 − 0.30)    where  t ∈ [0,1]\n"
    "\n"
    "t = clip( (cobalt_price − 0.80) / (1.50 − 0.80), 0, 1 )\n"
    "\n"
    "LFP_share(t+1) = LFP_share(t) + speed × (LFP_target − LFP_share(t))\n"
    "\n"
    "speed = 0.0045 (shifting toward LFP),  0.0015 (drifting back to NMC)\n"
    "LFP_share bounded to [0.15, 0.92]")

add_heading(doc, "F3 — Cell Capacity Investment Cycle (Reinforcing + Balancing)", 3)
add_para(doc,
    "High capacity utilisation triggers investment in new gigafactory capacity. "
    "A 104-week construction pipeline represents the real-world ~2-year build time "
    "from groundbreaking to first production. Once online, new capacity reduces "
    "utilisation, damping the investment signal.")

add_para(doc, "Causal chain:", bold=True)
add_para(doc,
    "  Cell demand ↑  →  utilisation ↑ (> 85%)  →  investment initiated  →  "
    "capacity WIP ↑  →  [104-week delay]  →  capacity online ↑  →  "
    "utilisation ↓  →  investment slows")

add_equation_box(doc,
    "excess_util = max(0,  utilisation − 0.85)\n"
    "investment_rate = capacity × (excess_util / 0.15) × 0.20 / 52   [GWh/wk]\n"
    "\n"
    "WIP(t+1) = WIP(t) + investment_rate − completion_rate\n"
    "completion_rate = WIP(t) / 104              [first-order delay]\n"
    "\n"
    "capacity(t+1) = capacity(t) + completion_rate − depreciation\n"
    "depreciation = capacity(t) / (52 × 15)      [15-year economic life]")

add_heading(doc, "F4 — Demand–Adoption Loop (Reinforcing + Balancing)", 3)
add_para(doc,
    "EV demand grows at a policy-driven baseline rate (29%/yr, IEA GEO 2024). "
    "Battery price signals modulate only the growth increment, not the base level "
    "of demand, consistent with empirical evidence that short-run price elasticity "
    "of EV demand is low while long-run growth-rate elasticity is moderate. "
    "This prevents the price effect from compounding multiplicatively each week "
    "and causing demand collapse under prolonged price elevation.")

add_para(doc, "Causal chain:", bold=True)
add_para(doc,
    "  Battery price ↑  →  growth rate modulated downward  →  "
    "EV demand grows more slowly  →  cell outflows slow  →  "
    "mineral stocks partially recover  →  battery price eases")

add_equation_box(doc,
    "trend(t+1) = trend(t) × (1 + weekly_growth_rate)     [price-independent]\n"
    "\n"
    "price_level = 1.0 + (−0.30) × (price_signal − 1.0)  [clamped to 0.50–1.20]\n"
    "\n"
    "weekly_demand_gwh = trend(t+1) × price_level\n"
    "\n"
    "EV demand grows at:  +29%/yr baseline\n"
    "                     +33%/yr (China),  +35%/yr (USA),  +28%/yr (UK)")

add_heading(doc, "F5 — Bullwhip Amplification Tracking", 3)
add_para(doc,
    "The bullwhip effect — upstream demand amplification — is tracked by comparing "
    "total Tier-1 component order rates against total component demand at the "
    "Tier-3/OEM boundary. A bullwhip_index > 1 indicates that upstream order volumes "
    "exceed downstream consumption, a key indicator of supply-chain instability. "
    "The index is smoothed with an exponentially weighted moving average (EWMA) "
    "to reduce week-to-week noise.")

add_equation_box(doc,
    "raw_bullwhip = total_component_order_rate / (4 × vehicle_demand)\n"
    "  (÷4 because 4 subsystem types per vehicle)\n"
    "\n"
    "bullwhip(t+1) = (1 − 0.10) × bullwhip(t)  +  0.10 × raw_bullwhip(t)\n"
    "  (EWMA with smoothing weight 0.10)")

# ── 3.3 Causal Loop Diagram ───────────────────────────────────────────────────
add_heading(doc, "3.3  Causal Loop Diagram", 2)

add_para(doc,
    "The Causal Loop Diagram (CLD) below describes all endogenous feedback "
    "relationships in the model. Arrows indicate causal influence; "
    "'+' means same direction (increase → increase), "
    "'−' means opposite direction (increase → decrease). "
    "Loop polarity: 'B' = balancing (negative feedback, stabilising), "
    "'R' = reinforcing (positive feedback, amplifying).")

doc.add_paragraph()
add_para(doc, "Top-level loop structure:", bold=True)
add_table(doc,
    ["Loop", "Type", "Key Variables", "Mechanism"],
    [
        ["F1", "Balancing (B1)",
         "Mineral stock → Price → Demand → Mineral outflow",
         "Shortage drives prices up; higher prices slow demand growth, reducing outflows"],
        ["F2", "Balancing (B2)",
         "Cobalt price → LFP share → Cobalt consumption → Cobalt stock",
         "Expensive cobalt accelerates LFP adoption; LFP uses no cobalt"],
        ["F3a", "Reinforcing (R1) short-run",
         "Cell demand → Utilisation → Investment → Capacity WIP",
         "Higher demand drives capacity investment"],
        ["F3b", "Balancing (B3) long-run",
         "Capacity WIP → Capacity online → Utilisation ↓",
         "New capacity comes online, reducing utilisation and investment signal"],
        ["F4", "Balancing (B4)",
         "Battery price → EV demand growth rate",
         "Higher battery prices moderate EV adoption growth"],
        ["F5", "Amplifying (R2)",
         "Shortfall → Order amplification → Bullwhip",
         "Inventory shortfalls trigger over-ordering; amplification measured vs. actual demand"],
    ],
    col_widths=[0.5, 1.3, 2.4, 2.8]
)

doc.add_paragraph()
add_para(doc, "Detailed causal chain notation:", bold=True)
add_para(doc, "F1 — Price Formation Loop:")
add_equation_box(doc,
    "[+] Mineral demand  →(+)→  Stock depletion\n"
    "[−] Stock depletion →(−)→  Availability fraction\n"
    "[+] Low availability→(+)→  Commodity price\n"
    "[+] High price      →(+)→  Battery pack price\n"
    "[−] High pack price →(−)→  EV demand growth\n"
    "[−] Lower demand    →(−)→  Mineral outflow\n"
    "[+] Lower outflow   →(+)→  Mineral stock recovery\n"
    "==> LOOP POLARITY: BALANCING (B1)")

doc.add_paragraph()
add_para(doc, "F2 — Chemistry Substitution Loop:")
add_equation_box(doc,
    "[+] Cobalt price    →(+)→  LFP adoption incentive\n"
    "[+] LFP incentive   →(+)→  LFP share target\n"
    "[+] LFP target      →(+)→  Actual LFP share (via first-order convergence)\n"
    "[−] Higher LFP share→(−)→  Cobalt consumption\n"
    "[+] Less cobalt use →(+)→  Cobalt stock\n"
    "[−] Higher stock    →(−)→  Cobalt price\n"
    "==> LOOP POLARITY: BALANCING (B2)")

doc.add_paragraph()
add_para(doc, "F3 — Capacity Investment Cycle:")
add_equation_box(doc,
    "[+] EV demand growth →(+)→  Cell production need\n"
    "[+] Production need  →(+)→  Capacity utilisation\n"
    "[+] High utilisation →(+)→  Investment trigger (> 85%)\n"
    "[+] Investment       →(+)→  Capacity WIP\n"
    "[+] WIP (104-wk lag) →(+)→  New capacity\n"
    "[−] New capacity     →(−)→  Utilisation\n"
    "[−] Lower utilisation→(−)→  Investment rate\n"
    "==> SHORT-RUN: REINFORCING (R1)  |  LONG-RUN: BALANCING (B3)")

doc.add_paragraph()
add_para(doc, "Cross-loop interactions:", bold=True)
add_bullet(doc,
    "F1 × F2: The cobalt price (F1 output) is the input trigger for chemistry "
    "substitution (F2), creating a direct coupling between the price loop and "
    "the chemistry loop. When cobalt price rises above 1.5× baseline, "
    "both loops operate simultaneously.")
add_bullet(doc,
    "F2 × F1: Increased LFP share (F2 output) reduces the cobalt weight in the "
    "battery price index (BOM calculation), slightly dampening the F1 price signal. "
    "This is an implicit cross-loop feedback.")
add_bullet(doc,
    "F3 × F4: Capacity investment (F3) increases cell supply, which reduces "
    "battery prices (F1) and therefore stimulates demand growth (F4), "
    "creating a reinforcing cross-loop between supply expansion and demand.")

# ── 3.4 Stock-Flow Equations ──────────────────────────────────────────────────
add_heading(doc, "3.4  Stock–Flow Equations (Euler Integration)", 2)

add_para(doc,
    "All stocks are updated by Euler integration with dt = 1 week. "
    "The following equations define the complete SD model dynamics.")

add_para(doc, "Mineral inventory stocks (Tier 1):", bold=True)
add_equation_box(doc,
    "S_mineral(t+1) = S_mineral(t) + InFlow(t) − OutFlow(t)\n"
    "\n"
    "InFlow_mineral(t) = ABM_supply_fraction × BASELINE_WK[mineral]\n"
    "                    × supply_expansion_scale(t)\n"
    "\n"
    "supply_expansion_scale(t+1) = supply_expansion_scale(t)\n"
    "                              × (1 + MINERAL_SUPPLY_GROWTH_WK[mineral])\n"
    "\n"
    "OutFlow_mineral(t) = ABM_consumption_fraction × BASELINE_WK[mineral]\n"
    "\n"
    "Stock cap: S_mineral ≤ 4 × target_weeks × BASELINE_WK[mineral]")

add_para(doc, "Commodity price dynamics (Tier 1):", bold=True)
add_equation_box(doc,
    "p_mineral(t+1) = p_mineral(t)\n"
    "               + PRICE_ADJ_SPEED × (1/availability(t) − p_mineral(t))\n"
    "\n"
    "availability(t) = S_mineral(t) / (target_weeks × BASELINE_WK[mineral])\n"
    "p_mineral bounded to [PRICE_FLOOR=0.10, PRICE_CEIL=6.00]")

add_para(doc, "Battery price signal (Tier 2 input):", bold=True)
add_equation_box(doc,
    "P_battery = [0.30 × p_Li  +  0.15×(1−LFP_share) × p_Co\n"
    "            + 0.15 × p_Graphite  +  0.05 × p_REE\n"
    "            + 0.10 × p_SiC  +  (0.25 + 0.15×LFP_share)] / total_weight\n"
    "\n"
    "Note: cobalt weight is blended by chemistry mix;\n"
    "      LFP cells have zero cobalt content.")

add_para(doc, "Cell inventory (Tier 2):", bold=True)
add_equation_box(doc,
    "S_cells(t+1) = S_cells(t) + cells_produced(t) − cells_consumed(t)\n"
    "\n"
    "cells_produced = Σ output_gwh across all CellManufacturerAgents\n"
    "cells_consumed = battery_pack_agent.output_k × kWh_per_vehicle / 1000")

add_para(doc, "Component stocks (Tier 3):", bold=True)
add_equation_box(doc,
    "For c ∈ {packs, inverters, motors, harness}:\n"
    "S_c(t+1) = S_c(t) + T1_output_c(t) − OEM_production(t)")

add_para(doc, "EV demand (Tier 4 — SD stock):", bold=True)
add_equation_box(doc,
    "growth_increment(t) = ev_demand(t) × EV_DEMAND_GROWTH_WK × price_effect(t)\n"
    "ev_demand(t+1) = ev_demand(t) + growth_increment(t)\n"
    "\n"
    "price_effect = clip(1.0 + (−0.30) × (P_battery − 1.0),  0.60, 1.20)\n"
    "\n"
    "Soft coupling (10% weekly adjustment) to ABM market-agent aggregate:\n"
    "ev_demand(t+1) = 0.90 × ev_demand(t+1) + 0.10 × (market_gwh_wk × 52)")

# ── 3.5 Parameterisation ─────────────────────────────────────────────────────
add_heading(doc, "3.5  SD Parameterisation", 2)
add_table(doc,
    ["Parameter", "Value", "Source"],
    [
        ["PRICE_ADJ_SPEED", "0.05 /wk",
         "Calibrated to Li₂CO₃ price doubling in ~20 weeks (2021 surge); Benchmark Mineral Intelligence"],
        ["PRICE_FLOOR", "0.10", "Numerical stability; extreme supply glut"],
        ["PRICE_CEILING", "6.00 × baseline", "Extreme stress cap (DRC total shutdown scenario)"],
        ["CAPEX_TRIGGER_UTIL", "0.85", "Industry rule of thumb; BNEF Gigafactory Tracker 2023"],
        ["CELL_CAPACITY_BUILD_WK", "104 weeks", "BNEF 2023; Tesla/CATL/LG ES new plant build times"],
        ["CHEM_SHIFT_SPEED (upward)", "0.0045 /wk", "Full LFP shift over ~3 years; CATL 2022–24 transition rate"],
        ["CHEM_SHIFT_SPEED (downward)", "0.0015 /wk", "Slower: NMC lines not abandoned quickly"],
        ["COBALT_SWITCH_LOW", "0.80 × baseline", "Below: weak NMC→LFP pressure"],
        ["COBALT_SWITCH_HIGH", "1.50 × baseline", "Above: full LFP incentive activated"],
        ["LFP_SHARE_MIN", "0.15", "Technological floor (NCA/NMC premium segment)"],
        ["LFP_SHARE_MAX", "0.92", "Practical ceiling (some NMC for high energy density)"],
        ["EV_DEMAND_GROWTH_WK", "0.491% /wk", "29%/yr IEA GEO 2024 baseline projection"],
        ["BULLWHIP_SMOOTH", "0.10", "EWMA weight; standard in supply-chain literature"],
        ["BULLWHIP_FACTOR", "1.25", "Tier-1 over-ordering multiplier; Lee et al. 1997"],
    ],
    col_widths=[2.0, 1.4, 3.6]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. AGENT-BASED MODEL
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4.  Agent-Based Model (ABM) Layer", 1)

add_para(doc,
    "The ABM layer models the heterogeneous behaviour of 40 individual firm-level "
    "agents across five types. Each agent carries a financial profile calibrated "
    "from listed-company data, which modifies four behavioural multipliers: "
    "recovery_multiplier, inventory_multiplier, growth_multiplier, and shock_absorption.")

add_heading(doc, "4.1  MineralSupplierAgent  (Tier 1 — 14 agents)", 2)

add_para(doc,
    "Models a country-level mineral production source (not an individual mine). "
    "Fourteen agents cover five minerals across different geographies.")

add_table(doc,
    ["Agent ID", "Mineral", "Country", "Global Share"],
    [
        ["lithium_aus",   "Lithium", "Australia",  "46%"],
        ["lithium_chl",   "Lithium", "Chile",       "30%"],
        ["lithium_chn",   "Lithium", "China",       "14%"],
        ["lithium_other", "Lithium", "Others",      "10%"],
        ["cobalt_drc",    "Cobalt",  "DRC",         "70%"],
        ["cobalt_other",  "Cobalt",  "Others",      "30%"],
        ["graphite_chn",  "Graphite","China",       "79%"],
        ["graphite_other","Graphite","Others",      "21%"],
        ["ree_chn",       "REE",     "China",       "85%"],
        ["ree_other",     "REE",     "Others",      "15%"],
        ["sic_wolfspeed", "SiC",     "USA",         "30%"],
        ["sic_coherent",  "SiC",     "USA",         "20%"],
        ["sic_china",     "SiC",     "China",       "18%"],
        ["sic_other",     "SiC",     "EU/Others",   "32%"],
    ],
    col_widths=[1.4, 1.1, 1.4, 1.3]
)

doc.add_paragraph()
add_para(doc, "Decision rule:", bold=True)
add_equation_box(doc,
    "output_fraction(t) = shock_multiplier(t)\n"
    "\n"
    "If shocked:\n"
    "  shock_multiplier = max(0,  1 − effective_severity)\n"
    "  effective_severity = shock_severity × (1 − shock_absorption)\n"
    "\n"
    "Recovery (when shock resolved):\n"
    "  shock_multiplier(t+1) = min(1.0,\n"
    "      shock_multiplier(t) + recovery_rate_wk × recovery_multiplier)\n"
    "\n"
    "SD inflow contribution:\n"
    "  weekly_supply_contribution = global_share × output_fraction")

add_heading(doc, "4.2  CellManufacturerAgent  (Tier 2 — 9 agents)", 2)

add_para(doc,
    "Models a battery cell manufacturer. The nine agents cover the major global "
    "cell makers plus an aggregate 'others' group.")

add_table(doc,
    ["Agent", "Country", "Market Share", "Capacity (GWh/yr)", "LFP%", "NMC%"],
    [
        ["CATL",         "China",         "37.0%", "304.1", "45%", "55%"],
        ["BYD Cells",    "China",         "14.0%", "115.1", "90%", "10%"],
        ["LG Energy Sol.","South Korea",  "13.0%", "106.9",  "5%", "95%"],
        ["Panasonic",    "Japan",          "7.0%",  "57.5",  "0%","100%"],
        ["Samsung SDI",  "South Korea",    "6.0%",  "49.3",  "0%","100%"],
        ["SK On",        "South Korea",    "5.0%",  "41.1",  "0%","100%"],
        ["CALB",         "China",          "5.0%",  "41.1", "80%", "20%"],
        ["AESC (UK)",    "UK",             "0.2%",   "1.6",  "0%","100%"],
        ["Others",       "Mixed",         "12.8%", "105.2", "50%", "50%"],
    ],
    col_widths=[1.5, 1.2, 1.1, 1.5, 0.6, 0.6]
)

doc.add_paragraph()
add_para(doc, "Leontief production constraint:", bold=True)
add_equation_box(doc,
    "cobalt_price_adj = clip((cobalt_price − 1.0) × 0.15,  0, 1)\n"
    "effective_NMC = nmc_fraction × (1 − cobalt_price_adj)\n"
    "effective_LFP = 1 − effective_NMC\n"
    "\n"
    "cobalt_constraint = effective_LFP + effective_NMC × availability[cobalt]\n"
    "\n"
    "leontief = min(availability[lithium],\n"
    "              availability[graphite],\n"
    "              cobalt_constraint)\n"
    "\n"
    "max_production = weekly_capacity × shock_multiplier × leontief\n"
    "\n"
    "Capacity growth each week:\n"
    "weekly_capacity(t+1) = weekly_capacity(t) × (1 + 0.491%/wk × growth_multiplier)")

add_para(doc, "Inventory and order policy:", bold=True)
add_equation_box(doc,
    "target_inventory = weekly_capacity × safety_stock_weeks × inventory_multiplier\n"
    "gap = max(0,  target_inventory − current_inventory)\n"
    "desired_output = weekly_capacity + gap / 4.0\n"
    "actual_output  = min(desired_output,  max_production)\n"
    "\n"
    "inventory(t+1) = inventory(t) + actual_output − fulfilled_demand")

add_heading(doc, "4.3  Tier1SupplierAgent  (Tier 3 — 4 agents)", 2)

add_para(doc,
    "Models a sub-system supplier. Four agents cover the four critical EV sub-assemblies. "
    "Each agent maintains an explicit lead-time order pipeline.")

add_table(doc,
    ["Agent", "Component", "Key Input", "Input Dep.", "Lead Time", "Safety Stock"],
    [
        ["t1_battery_pack","Battery pack",  "Cells",     "100%", "4 wks",  "3 wks"],
        ["t1_inverter",    "Inverter",      "SiC wafer", " 45%", "16 wks", "8 wks"],
        ["t1_motor",       "Motor",         "REE",       " 82%", "12 wks", "6 wks"],
        ["t1_harness",     "Wiring harness","Copper",    "  0%", " 6 wks", "2 wks"],
    ],
    col_widths=[1.4, 1.3, 1.1, 0.9, 1.0, 1.1]
)

doc.add_paragraph()
add_para(doc, "Decision rule — order-up-to with bullwhip:", bold=True)
add_equation_box(doc,
    "Partial input constraint (allows non-critical portion to run):\n"
    "  constraint = (1 − input_dependency) + input_dependency × availability[key_input]\n"
    "\n"
    "Order-up-to policy:\n"
    "  inventory_position = on_hand_inventory + in_transit(pipeline)\n"
    "  shortfall = max(0,  target_inventory − inventory_position)\n"
    "  order = weekly_target + BULLWHIP_FACTOR × shortfall\n"
    "\n"
    "Dual sourcing (activates when inventory < 20% of target):\n"
    "  order × = 1.20  (20% cost premium for expedited sourcing)\n"
    "\n"
    "Pipeline: FIFO queue of length lead_time_weeks\n"
    "  pipeline.append(order)   [place order]\n"
    "  delivery = pipeline.pop(0)   [receive oldest order]")

add_heading(doc, "4.4  OEMAgent  (Tier 4 — 7 agents)", 2)

add_para(doc,
    "Models a vehicle manufacturer group. Seven agents cover major regional "
    "OEM groups. Production is governed by a strict Leontief constraint across "
    "all four sub-system inputs.")

add_table(doc,
    ["Agent", "Region", "Annual Target (k)", "Vertical Integration", "Safety Stock"],
    [
        ["byd_oem",          "China",  "1,575",  "95% (Blade cells, motors, chips in-house)", "3 wks"],
        ["other_chinese_oem","China",  "6,825",  "50% (CATL-reliant)",                        "4 wks"],
        ["uk_oem",           "UK",       "175",  " 5% (no domestic cell capacity)",            "5 wks"],
        ["us_oem",           "USA",    "1,820",  "40% (Tesla in-house + GM Ultium)",           "5 wks"],
        ["german_oem",       "Europe", "1,505",  "20% (heavy Tier-1 reliance)",                "6 wks"],
        ["korean_oem",       "Korea",  "1,120",  "45%",                                        "5 wks"],
        ["japanese_oem",     "Japan",    "980",  "30% (Panasonic partnership)",                "6 wks"],
    ],
    col_widths=[1.6, 0.8, 1.3, 2.4, 0.9]
)

doc.add_paragraph()
add_para(doc, "Production function:", bold=True)
add_equation_box(doc,
    "For each OEM, at each step:\n"
    "\n"
    "  1. Receive deliveries from all 4 Tier-1 agents (proportional to OEM market share)\n"
    "  2. Leontief constraint:\n"
    "       producible = min(inv_packs, inv_inverters, inv_motors, inv_harness)\n"
    "  3. Target production:\n"
    "       target = min(weekly_demand + 0.15 × backlog,  weekly_target × 1.10)\n"
    "  4. Actual production:\n"
    "       production_k = min(target, producible) × shock_multiplier\n"
    "  5. Consume components:\n"
    "       inv[c](t+1) = inv[c](t) − production_k   for c ∈ {packs,inverters,motors,harness}\n"
    "  6. Backlog update:\n"
    "       backlog(t+1) = max(0, backlog(t) + shortfall − 0.5 × surplus)\n"
    "  7. Halt detection:\n"
    "       if production_k < 0.10 × weekly_target: halt_weeks += 1")

add_heading(doc, "4.5  MarketAgent  (Demand — 6 agents)", 2)

add_para(doc,
    "Generates exogenous EV battery demand for each regional market. "
    "Demand follows a price-independent trend growth with a non-compounding "
    "price-level adjustment applied each week.")

add_table(doc,
    ["Region", "GWh 2023", "YoY Growth", "Avg kWh/Vehicle", "Note"],
    [
        ["China",  "493 GWh", "+33%/yr", "62 kWh", "IEA GEO 2024"],
        ["Europe", "127 GWh", "+11%/yr", "68 kWh", "Excl. UK"],
        ["UK",      "20 GWh", "+28%/yr", "65 kWh", "ZEV mandate 2024; SMMT 2024"],
        ["USA",    "104 GWh", "+35%/yr", "85 kWh", "IEA GEO 2024"],
        ["Japan",   "14 GWh", " +6%/yr", "48 kWh", "Japan + Korea combined"],
        ["ROW",     "64 GWh", "+42%/yr", "55 kWh", "SE Asia, LatAm, Middle East"],
    ],
    col_widths=[0.9, 0.9, 1.0, 1.4, 2.8]
)

doc.add_paragraph()
add_para(doc, "Decision rule:", bold=True)
add_equation_box(doc,
    "Weekly growth:\n"
    "  _trend_demand(t+1) = _trend_demand(t) × (1 + weekly_growth_rate)\n"
    "  weekly_growth_rate = (1 + yoy_growth)^(1/52) − 1\n"
    "\n"
    "Price-level adjustment (non-compounding):\n"
    "  price_signal  = SD_model.get_price_signal()\n"
    "  price_level   = clip(1.0 + (−0.30) × (price_signal − 1.0),  0.50, 1.20)\n"
    "\n"
    "Observed demand:\n"
    "  weekly_demand_gwh(t) = _trend_demand(t+1) × price_level(t)\n"
    "\n"
    "  NOTE: price_level is applied to the TREND (not the previous week's demand),\n"
    "  preventing the price effect from compounding multiplicatively over time.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. HYBRID COUPLING
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5.  Hybrid Coupling: ABM ↔ SD", 1)

add_para(doc,
    "The hybrid model architecture ensures that the SD and ABM layers remain "
    "tightly coupled without either layer dominating the other. Coupling "
    "operates in both directions each week.")

add_heading(doc, "5.1  SD → ABM (Before Agent Steps)", 2)
add_table(doc,
    ["Signal", "SD Source", "ABM Recipient", "Effect"],
    [
        ["input_fractions[mineral]", "stocks[mineral] / target",
         "CellManufacturerAgent, Tier1SupplierAgent",
         "Leontief constraint on production"],
        ["prices[cobalt]", "_step_prices()",
         "CellManufacturerAgent",
         "Reduces effective NMC fraction (chemistry adaptation)"],
        ["price_signal", "_battery_price_signal()",
         "MarketAgent",
         "Adjusts weekly demand via price-level factor"],
        ["lfp_share", "_step_chemistry_mix()",
         "Hybrid _collect_flows()",
         "Modulates cobalt outflow calculation"],
    ],
    col_widths=[1.5, 1.5, 1.8, 2.2]
)

add_heading(doc, "5.2  ABM → SD (After Agent Steps)", 2)
add_table(doc,
    ["Flow Key", "ABM Source", "SD Destination", "Unit"],
    [
        ["lithium_in / out", "MineralSupplierAgents", "stocks[lithium]", "fraction of baseline"],
        ["cobalt_in / out",  "MineralSupplierAgents + chemistry mix", "stocks[cobalt]", "fraction of baseline"],
        ["graphite_in / out","MineralSupplierAgents", "stocks[graphite]", "fraction of baseline"],
        ["ree_in / out",     "MineralSupplierAgents + motor output", "stocks[ree]", "fraction of baseline"],
        ["sic_wafer_in/out", "MineralSupplierAgents + inverter output", "stocks[sic_wafer]", "fraction of baseline"],
        ["cells_in / out",   "CellManufacturerAgents / pack agent", "stocks[cells]", "GWh"],
        ["packs/inv/motor/harness _in/_out", "Tier1SupplierAgents / OEM production", "stocks[packs…]", "k veh-equiv"],
        ["cell_capacity_gwh_yr", "sum of cell agent capacities × 52", "cell_capacity (sync)", "GWh/yr"],
        ["total_oem_prod_k", "sum OEM production", "OEM backlog update", "k vehicles/wk"],
        ["total_demand_k",   "sum market demand (k veh)", "Bullwhip + backlog", "k vehicles/wk"],
        ["order_rate_k",     "sum of Tier-1 pipeline additions", "Bullwhip numerator", "k units/wk"],
    ],
    col_widths=[1.8, 1.8, 1.6, 1.2]
)

add_heading(doc, "5.3  Step Execution Order", 2)
add_para(doc,
    "The strict ordering ensures that each agent uses the most recent SD state "
    "and that the SD model receives the fully updated agent outputs before "
    "advancing its own stocks:")
add_equation_box(doc,
    "WEEK t:\n"
    "  1.  apply_shocks()                         [inject exogenous disruptions]\n"
    "  2.  sd.compute_input_fractions()            [SD → ABM: availability signals]\n"
    "  3.  mineral_agents.step()                   [ABM Tier 1]\n"
    "  4.  cell_agents.step()                      [ABM Tier 2]\n"
    "  5.  tier1_agents.step()                     [ABM Tier 3]\n"
    "  6.  oem_agents.step()                       [ABM Tier 4]\n"
    "  7.  market_agents.step()                    [ABM Demand]\n"
    "  8.  flows = _collect_flows()                [aggregate ABM outputs]\n"
    "  9.  sd.update(flows)                        [ABM → SD: update all stocks]\n"
    "       ↳ _step_prices()                       [F1: price formation]\n"
    "       ↳ _step_chemistry_mix()                [F2: LFP substitution]\n"
    "       ↳ _step_cell_capacity(flows)           [F3: capacity investment]\n"
    "       ↳ _step_demand_and_backlog(flows)      [F4: demand dynamics]\n"
    "       ↳ _step_bullwhip(flows)               [F5: order amplification]\n"
    " 10.  sd.record()                             [snapshot SD state]\n"
    " 11.  _record_metrics()                       [snapshot all metrics → DataFrame]")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 6. FINANCIAL CALIBRATION
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "6.  Financial Calibration from Listed Companies", 1)

add_para(doc,
    "A core contribution of this model is the calibration of agent behavioural "
    "parameters from real listed-company financial data. For each model agent, "
    "a set of peer companies is identified from the corresponding supply-chain tier. "
    "Five years of Yahoo Finance fundamentals are retrieved and used to compute "
    "four behavioural multipliers.")

add_heading(doc, "6.1  Financial Profile Multipliers", 2)
add_table(doc,
    ["Multiplier", "Derived From", "Effect on Agent Behaviour", "Default"],
    [
        ["recovery_multiplier",  "Revenue volatility / CAGR",
         "Scales the weekly recovery rate after a shock; financially stronger firms recover faster",
         "1.0"],
        ["inventory_multiplier", "Days Inventory Outstanding (DIO)",
         "Scales safety-stock target; firms with longer cash cycles hold more inventory",
         "1.0"],
        ["growth_multiplier",    "Revenue 5-yr CAGR",
         "Scales capacity growth rate; high-growth firms expand capacity faster",
         "1.0"],
        ["shock_absorption",     "Net margin / balance-sheet strength proxy",
         "Reduces effective shock severity; firms with stronger finances absorb disruptions better",
         "0.0"],
    ],
    col_widths=[1.6, 1.8, 2.8, 0.8]
)

add_heading(doc, "6.2  Peer Company Groups by Tier", 2)
add_table(doc,
    ["Tier", "Agent Group", "Example Listed Peers"],
    [
        ["Tier 1 — Minerals",
         "Lithium miners, cobalt producers, graphite producers, REE, SiC",
         "Albemarle (ALB), SQM, Ganfeng Lithium, Glencore (GLEN.L), CMOC (3993.HK), "
         "MP Materials (MP), Lynas Rare Earths (LYC.AX), Wolfspeed (WOLF), Coherent (COHR)"],
        ["Tier 2 — Cells",
         "Battery cell manufacturers",
         "CATL (300750.SZ), BYD (002594.SZ), LG Energy Solution (373220.KS), "
         "Panasonic (6752.T), Samsung SDI (006400.KS), SK On (096770.KS), CALB (3931.HK)"],
        ["Tier 3 — Subsystems",
         "Pack, inverter, motor, harness suppliers",
         "BorgWarner (BWA), Infineon (IFX.DE), ON Semiconductor (ON), STMicro (STM), "
         "Nidec (6594.T), Denso (6902.T), Aptiv (APTV), TE Connectivity (TEL)"],
        ["Tier 4 — OEM",
         "Vehicle manufacturers by region",
         "Tesla (TSLA), GM (GM), Ford (F), Volkswagen (VOW3.DE), BMW (BMW.DE), "
         "Hyundai (005380.KS), Toyota (7203.T), BYD (002594.SZ), NIO (NIO), Li Auto (LI)"],
    ],
    col_widths=[1.3, 1.8, 3.9]
)

add_para(doc, "")
add_para(doc,
    "In total, 61 peer companies with valid tickers were identified across 34 agent "
    "groups. Financial data was successfully retrieved for 60 of these companies "
    "(Wolfspeed returned no records due to Chapter 11 bankruptcy filing in 2025; "
    "Coherent acts as fallback). Coverage is 100% for 33 of 34 agent groups.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 7. SCENARIO SHOCKS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "7.  Scenario Shocks", 1)

add_para(doc,
    "The model supports structured scenario shocks applied to any agent. "
    "Shocks are defined as a list of shock dictionaries in the scenario configuration, "
    "specifying the target agent, start week, end week, and severity [0, 1].")

add_table(doc,
    ["Scenario", "Target Agent", "Severity", "Duration", "Mechanism"],
    [
        ["DRC Cobalt Disruption",
         "cobalt_drc",
         "0.70 (70% output loss)",
         "Weeks 26–78 (12 months)",
         "Political instability / export ban. Cobalt stock depletes → price rises → "
         "LFP substitution accelerates. Modelled qualitative result: cobalt price +86%, "
         "LFP share rises from 40% to 45%."],
        ["China REE Export Restriction",
         "ree_chn",
         "0.80 (80% output loss)",
         "Weeks 26–52 (6 months)",
         "Export quota reduction. REE price rises 2.5–4.2×. Motor production "
         "constrained by REE shortage. OEM production reduced."],
        ["China Graphite Disruption",
         "graphite_chn",
         "0.60 (60% output loss)",
         "Weeks 26–60",
         "China controls 79% of global graphite supply. Cell production constrained "
         "as graphite anode shortage develops."],
        ["UK Supply Chain Friction",
         "uk_oem",
         "0.15 (15% throughput loss)",
         "Weeks 1–260 (persistent)",
         "Post-Brexit rules-of-origin compliance costs and border friction. "
         "UK OEM operates at 85% of baseline capacity throughout."],
        ["SiC Wafer Shortage",
         "sic_wolfspeed",
         "0.50 (50% output loss)",
         "Weeks 26–52",
         "Wolfspeed capacity constraints. Inverter production partially constrained "
         "(45% SiC dependency); Si-IGBT inverters continue unaffected."],
    ],
    col_widths=[1.5, 1.3, 1.0, 1.2, 3.0]
)

doc.add_paragraph()
add_para(doc, "Shock application mechanism:", bold=True)
add_equation_box(doc,
    "effective_severity = shock_severity × (1 − agent.financial_profile.shock_absorption)\n"
    "\n"
    "agent.shock_multiplier = max(0,  1 − effective_severity)\n"
    "\n"
    "Recovery (once end_week is reached):\n"
    "  each week: shock_multiplier += recovery_rate_wk × recovery_multiplier\n"
    "  until shock_multiplier = 1.0\n"
    "\n"
    "Example: DRC cobalt (severity=0.70, absorption=0.05)\n"
    "  effective_severity = 0.70 × 0.95 = 0.665\n"
    "  shock_multiplier = 0.335  (cobalt output drops to 33.5% of normal)")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 8. PARAMETER SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "8.  Complete Parameter Summary", 1)

add_heading(doc, "8.1  Simulation Control", 2)
add_table(doc,
    ["Parameter", "Value", "Description"],
    [
        ["dt",       "1 week",    "Time step"],
        ["n_weeks",  "260",       "Simulation horizon (5 years)"],
        ["seed",     "42",        "Random number seed for reproducibility"],
        ["warm_up",  "4 weeks",   "Initial stabilisation period before metrics"],
    ],
    col_widths=[1.5, 1.0, 4.5]
)

add_heading(doc, "8.2  Mineral Safety-Stock Targets", 2)
add_table(doc,
    ["Mineral", "Target (wks)", "Baseline (wks/unit)", "Rationale"],
    [
        ["Lithium",   "4",  "1.0 normalised", "IEA CMMR 2023 typical buffer"],
        ["Cobalt",    "6",  "1.0 normalised", "DRC political risk"],
        ["Graphite",  "4",  "1.0 normalised", "Standard anode buffer"],
        ["REE",       "8",  "1.0 normalised", "Slow processing; China dominant"],
        ["SiC wafer", "12", "1.0 normalised", "26-week lead time; strategic buffer"],
    ],
    col_widths=[1.2, 1.2, 1.8, 2.8]
)

add_heading(doc, "8.3  Mineral Supply Growth Rates", 2)
add_table(doc,
    ["Mineral", "Annual Growth", "Weekly Rate", "Source"],
    [
        ["Lithium",   "20%/yr", "0.352%/wk", "IEA CMMR 2023 — new capacity underway"],
        ["Cobalt",    " 5%/yr", "0.094%/wk", "CMOC/Glencore DRC expansions limited"],
        ["Graphite",  "15%/yr", "0.267%/wk", "Mozambique, Madagascar ramps"],
        ["REE",       " 8%/yr", "0.149%/wk", "MP Materials, Lynas; 3–5yr build times"],
        ["SiC wafer", "35%/yr", "0.586%/wk", "Wolfspeed, STM, Infineon fab ramps (BNEF 2023)"],
    ],
    col_widths=[1.2, 1.2, 1.2, 3.4]
)
add_para(doc, "Note: all supply growth rates are below EV demand growth (29%/yr), "
         "consistent with IEA and BNEF structural tightening forecasts.",
         italic=True)

add_heading(doc, "8.4  Key Structural Assumptions", 2)
add_bullet(doc,
    "One-for-one component ratios: each vehicle requires exactly one battery pack, "
    "one inverter, one motor, and one wiring harness set.")
add_bullet(doc,
    "Linear market share routing: cell demand is routed to cell makers and "
    "component deliveries to OEMs in proportion to their respective market shares. "
    "No switching behaviour is modelled (routing is static).")
add_bullet(doc,
    "Global homogeneous supply: mineral stocks are global aggregates. "
    "Geographic routing of specific minerals to specific cell makers is not modelled.")
add_bullet(doc,
    "No financial stress dynamics: agent financial profiles are fixed at "
    "calibrated values and do not evolve endogenously (e.g., balance sheets "
    "do not deteriorate under prolonged shock).")
add_bullet(doc,
    "No inventory sharing: agents cannot borrow from or share inventory "
    "with other agents at the same tier.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 9. REFERENCES
# ═══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "9.  References", 1)

refs = [
    ("Grimm, V., Berger, U., Bastiansen, F., et al. (2006).",
     "A standard protocol for describing individual-based and agent-based models. "
     "Ecological Modelling, 198(1-2), 115–126."),
    ("Grimm, V., Berger, U., DeAngelis, D.L., et al. (2010).",
     "The ODD protocol: A review and first update. Ecological Modelling, 221(23), 2760–2768."),
    ("Grimm, V., et al. (2020).",
     "The ODD Protocol for describing agent-based and other simulation models: A second update "
     "to improve clarity, replication, and structural realism. Journal of Artificial Societies "
     "and Social Simulation, 23(2), 7."),
    ("IEA (2024).",
     "Global EV Outlook 2024. International Energy Agency, Paris."),
    ("IEA (2023).",
     "Critical Minerals Market Review 2023. International Energy Agency, Paris."),
    ("USGS (2024).",
     "Mineral Commodity Summaries 2024. U.S. Geological Survey."),
    ("BNEF (2023).",
     "BloombergNEF Battery Price Survey 2023. Bloomberg New Energy Finance."),
    ("BNEF (2023).",
     "Battery Market Outlook 2023. Bloomberg New Energy Finance."),
    ("Yole Intelligence (2023).",
     "Power Electronics for EV/HEV — Market Report 2023."),
    ("Lee, H.L., Padmanabhan, V., & Whang, S. (1997).",
     "The bullwhip effect in supply chains. Sloan Management Review, 38(3), 93–102."),
    ("Benchmark Mineral Intelligence (2022).",
     "Lithium Carbonate Price Series 2021–2022."),
    ("SMMT (2024).",
     "SMMT New Car Registrations Data 2024. Society of Motor Manufacturers and Traders, UK."),
    ("UK DfT (2023).",
     "Zero Emission Vehicle (ZEV) Mandate: Consultation 2023. UK Department for Transport."),
    ("BYD Co. Ltd (2023).",
     "Annual Report 2023. BYD Company Limited."),
    ("CATL (2023).",
     "Annual Report 2023 (Chinese). Contemporary Amperex Technology Co. Limited."),
    ("Wolfspeed Inc. (2023).",
     "Annual Report on Form 10-K, FY2023. United States Securities and Exchange Commission."),
    ("Aptiv PLC (2023).",
     "Annual Report 2023. Aptiv PLC."),
]

for author, text in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.35)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run(author + " ")
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)

# ═══════════════════════════════════════════════════════════════════════════════
# Save
# ═══════════════════════════════════════════════════════════════════════════════
doc.save(OUTPUT_PATH)
print(f"Saved: {OUTPUT_PATH}")
